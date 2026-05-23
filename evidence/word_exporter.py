#!/usr/bin/env python3
"""
word_exporter.py - Export divergence cases to professional Microsoft Word reports.

v3.0: Reframed to "Governance Promise-Delivery Divergence Reconstruction Engine".
Added forensic branding and "Promise vs Outcome" layout.
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from evidence.evidence_gate import get_linked_evidence, validate_cases_for_report

class WordExporter:
    """Exports divergence cases to a professional Word document."""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _add_hyperlink(self, paragraph, text, url):
        """Adds a clickable hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0000FF')
        rPr.append(c)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
        return hyperlink

    def _add_evidence_details(self, doc, evidence_items):
        for evidence in evidence_items:
            info = doc.add_paragraph()
            info.add_run("EVIDENCE ID: ").bold = True
            info.add_run(str(evidence.get("evidence_id")))
            info.add_run("\nTITLE: ").bold = True
            info.add_run(str(evidence.get("title", "Unknown")))
            info.add_run("\nSOURCE TYPE: ").bold = True
            info.add_run(str(evidence.get("source_type", "Unknown")))
            info.add_run("\nPLATFORM: ").bold = True
            info.add_run(str(evidence.get("platform", "Unknown")))
            info.add_run("\nVERIFICATION STATUS: ").bold = True
            info.add_run(str(evidence.get("verification_status", "Unknown")))
            info.add_run("\nEVIDENCE STRENGTH: ").bold = True
            info.add_run(str(evidence.get("evidence_strength", "Unknown")))
            info.add_run("\nLINK: ").bold = True
            self._add_hyperlink(info, "View Evidence", evidence.get("url", "#"))

            timestamp_start = evidence.get("timestamp_start")
            timestamp_end = evidence.get("timestamp_end")
            if timestamp_start or timestamp_end:
                info.add_run("\nTIMESTAMP: ").bold = True
                info.add_run(" - ".join(str(t) for t in (timestamp_start, timestamp_end) if t))

    def generate_report(self, cases: List[Dict[str, Any]], filename: str = "governance_divergence_report.docx"):
        validate_cases_for_report(cases)

        doc = Document()
        
        # 1. Title Page
        self._add_title_page(doc, len(cases))
        doc.add_page_break()

        # 2. Executive Summary Table
        self._add_summary_table(doc, cases)
        doc.add_page_break()

        # 3. Case Sections
        for case in cases:
            self._add_case_section(doc, case)
            doc.add_page_break()

        # Save
        save_path = self.output_dir / filename
        doc.save(str(save_path))
        print(f"Report generated: {save_path}")
        return save_path

    def _add_title_page(self, doc, case_count):
        # Logo / Top Margin
        for _ in range(5): doc.add_paragraph()
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("DUMA BOKO")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("PROMISE DELIVERY & GOVERNANCE DIVERGENCE REPORT")
        run.font.size = Pt(18)
        run.font.bold = True

        for _ in range(3): doc.add_paragraph()

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f"REPORT DATE: {datetime.now().strftime('%B %d, %Y')}\n")
        run.add_text(f"TOTAL INVESTIGATION THEMES: {case_count}\n")
        run.add_text("CLASSIFICATION: HIGH-LEVEL FORENSIC AUDIT\n")
        run.add_text("STATUS: FINALIZED RECONSTRUCTION")
        run.font.size = Pt(12)

    def _add_summary_table(self, doc, cases):
        doc.add_heading("EXECUTIVE SUMMARY", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'INVESTIGATION THEME'
        hdr_cells[2].text = 'DIVERGENCE TYPE'
        hdr_cells[3].text = 'STRENGTH'

        for case in cases:
            row_cells = table.add_row().cells
            row_cells[0].text = str(case.get('case_id'))
            row_cells[1].text = str(case.get('topic', '')).replace('_', ' ').upper()
            row_cells[2].text = str(case.get('divergence_type', '')).replace('_', ' ')
            row_cells[3].text = str(case.get('evidence_strength', 'low')).upper()

    def _add_case_section(self, doc, case):
        case_id = case.get('case_id')
        topic = str(case.get('topic', '')).replace('_', ' ').upper()
        promise_evidence = get_linked_evidence(case, "promise")
        outcome_evidence = get_linked_evidence(case, "outcome_or_position")
        analysis_evidence = get_linked_evidence(case, "analysis")
        
        # Header
        heading = doc.add_heading(f"{case_id} — {topic}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Metadata
        p = doc.add_paragraph()
        run = p.add_run("DIVERGENCE TYPE: ")
        run.bold = True
        p.add_run(str(case.get('divergence_type', ''))).italic = True

        # Promise Block
        doc.add_heading("1. CAMPAIGN PROMISE (Earlier Position)", level=2)
        promise = case.get('promise', {})
        p = doc.add_paragraph(style='Intense Quote')
        p.add_run(f"\"{promise.get('quote', 'N/A')}\"")
        
        self._add_evidence_details(doc, promise_evidence)

        # Outcome Block
        doc.add_heading("2. GOVERNANCE OUTCOME (Later Position)", level=2)
        outcome = case.get('outcome_or_position', {})
        p = doc.add_paragraph(style='Intense Quote')
        p.add_run(f"\"{outcome.get('quote', 'N/A')}\"")
        
        self._add_evidence_details(doc, outcome_evidence)

        # Analysis
        doc.add_heading("3. RECONSTRUCTION ANALYSIS", level=2)
        doc.add_paragraph(str(case.get('analysis', 'No analysis available.')))
        evidence_note = doc.add_paragraph()
        evidence_note.add_run("ANALYSIS EVIDENCE IDS: ").bold = True
        evidence_note.add_run(
            ", ".join(str(evidence.get("evidence_id")) for evidence in analysis_evidence)
        )

        # Video Validation Flag
        doc.add_paragraph()
        v_para = doc.add_paragraph()
        run = v_para.add_run("✔ SOURCE-LINKED EVIDENCE ATTACHED")
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True

if __name__ == "__main__":
    import json
    with open('outputs/cases/divergence_cases.json', 'r') as f:
        data = json.load(f)
    exporter = WordExporter(Path("outputs/reports"))
    exporter.generate_report(data.get('cases', []), "DUMA_BOKO_DIVERGENCE_REPORT_FINAL.docx")
