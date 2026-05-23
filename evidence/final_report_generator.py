#!/usr/bin/env python3
"""
final_report_generator.py - Generates the high-level forensic audit report 
using user-provided ground truth text and verified video evidence.
"""

import os
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from evidence.evidence_gate import get_linked_evidence, validate_cases_for_report

class FinalReportGenerator:
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _add_hyperlink(self, paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0000FF')
        rPr.append(c)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_evidence_details(self, doc, evidence_items):
        for evidence in evidence_items:
            info = doc.add_paragraph()
            info.add_run("EVIDENCE ID: ").bold = True
            info.add_run(str(evidence.get("evidence_id")))
            info.add_run("\nTITLE: ").bold = True
            info.add_run(str(evidence.get("title", "Unknown")))
            info.add_run("\nSOURCE TYPE: ").bold = True
            info.add_run(str(evidence.get("source_type", "Unknown")))
            info.add_run("\nPLATFORM: ").bold = True
            info.add_run(str(evidence.get("platform", "Unknown")))
            info.add_run("\nVERIFICATION STATUS: ").bold = True
            info.add_run(str(evidence.get("verification_status", "Unknown")))
            info.add_run("\nEVIDENCE STRENGTH: ").bold = True
            info.add_run(str(evidence.get("evidence_strength", "Unknown")))
            info.add_run("\nLINK: ").bold = True
            self._add_hyperlink(info, "View Evidence", evidence.get("url", "#"))

            timestamp_start = evidence.get("timestamp_start")
            timestamp_end = evidence.get("timestamp_end")
            if timestamp_start or timestamp_end:
                info.add_run("\nTIMESTAMP: ").bold = True
                info.add_run(" - ".join(str(t) for t in (timestamp_start, timestamp_end) if t))

    def generate(self, cases: List[Dict[str, Any]], filename: str):
        validate_cases_for_report(cases)

        doc = Document()
        
        # Title Page
        for _ in range(3): doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run("DUMA BOKO")
        run.font.size = Pt(36); run.font.bold = True
        
        st = doc.add_paragraph()
        st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = st.add_run("PROMISE DELIVERY & GOVERNANCE DIVERGENCE REPORT")
        run.font.size = Pt(20); run.font.bold = True
        
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Report Date: {datetime.now().strftime('%B %d, %Y')} (Africa/Gaborone)\n")
        info.add_run("Status: High-Level Forensic Audit – Finalised Reconstruction")
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading("EXECUTIVE SUMMARY", level=1)
        doc.add_paragraph("This report evaluates major promises made by President Duma Boko and the Umbrella for Democratic Change (UDC) during the 2024 election campaign and assesses subsequent actions and outcomes using verifiable evidence.")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdrs = table.rows[0].cells
        hdrs[0].text = 'Case ID'; hdrs[1].text = 'Investigation Theme'; hdrs[2].text = 'Divergence Type'; hdrs[3].text = 'Evidence Strength'
        
        summary_data = [
            (
                str(case.get("case_id", "")),
                str(case.get("topic", "")).replace("_", " ").title(),
                str(case.get("divergence_type", "")).replace("_", " ").title(),
                str(case.get("evidence_strength", "")).title(),
            )
            for case in cases
        ]
        for cid, theme, dtype, strength in summary_data:
            row = table.add_row().cells
            row[0].text = cid; row[1].text = theme; row[2].text = dtype; row[3].text = strength
            
        doc.add_page_break()
        
        # Detailed Cases
        for case in cases:
            self._add_case_section(doc, case)
            doc.add_page_break()
            
        save_path = self.output_dir / filename
        doc.save(str(save_path))
        return save_path

    def _add_case_section(self, doc, case):
        topic = str(case.get('topic', '')).replace('_', ' ').title()
        doc.add_heading(f"{case['case_id']} — {topic}", level=1)
        doc.add_paragraph().add_run(f"Divergence Type: {case['divergence_type']}").bold = True
        doc.add_paragraph().add_run(
            f"Verification Status: {case.get('verification_status', 'unknown')}"
        ).bold = True
        
        doc.add_heading("1. Campaign Promise (Earlier Position)", level=2)
        promise = case.get('promise', {})
        doc.add_paragraph(str(promise.get('quote', 'N/A')), style='Intense Quote')
        self._add_evidence_details(doc, get_linked_evidence(case, "promise"))

        doc.add_heading("2. Governance Outcome (Later Position)", level=2)
        outcome = case.get('outcome_or_position', {})
        doc.add_paragraph(str(outcome.get('quote', 'N/A')), style='Intense Quote')
        self._add_evidence_details(doc, get_linked_evidence(case, "outcome_or_position"))

        doc.add_heading("3. Reconstruction Analysis", level=2)
        doc.add_paragraph(str(case.get('analysis', 'No analysis available.')))
        analysis_evidence = get_linked_evidence(case, "analysis")
        evidence_note = doc.add_paragraph()
        evidence_note.add_run("Analysis Evidence IDs: ").bold = True
        evidence_note.add_run(
            ", ".join(str(evidence.get("evidence_id")) for evidence in analysis_evidence)
        )
        
        p = doc.add_paragraph()
        run = p.add_run("✔ VALIDATED FORENSIC EVIDENCE LINKS ATTACHED")
        run.bold = True; run.font.color.rgb = RGBColor(0, 128, 0)
