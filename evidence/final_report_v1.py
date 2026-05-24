"""
final_report_v1.py - Controlled review-only report generation.
"""

from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List
import json

from docx import Document

from evidence.evidence_schema import save_json
from evidence.report_section_assembly import (
    DEFAULT_REPORT_SECTION_OUTPUT,
    AssembledReportSection,
    assemble_report_sections_fixture_only,
    validate_assembled_report_section,
)


DEFAULT_FINAL_REPORT_DIR = Path("outputs/final_report")
DEFAULT_FINAL_REPORT_PAYLOAD = DEFAULT_FINAL_REPORT_DIR / "final_report_payload.json"
DEFAULT_FINAL_REPORT_DOCX = (
    DEFAULT_FINAL_REPORT_DIR / "DUMA_BOKO_REVIEW_DRAFT_REPORT.docx"
)

FINAL_REPORT_STATUSES = {
    "review_draft",
    "blocked",
    "generation_failed",
}

FORBIDDEN_REVIEW_TEXT = (
    "validated public evidence",
    "final forensic report",
    "institution-ready",
    "proven corruption",
    "proven failure",
    "report_ready",
)


@dataclass
class FinalReportPayload:
    report_id: str
    title: str
    report_status: str
    generated_from: str
    sections: List[Dict[str, Any]]
    evidence_disclaimer: str
    methodology_note: str
    limitations_note: str
    generation_notes: str

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _as_dict(value: Any) -> Dict[str, Any]:
    if hasattr(value, "to_dict"):
        return value.to_dict()
    if isinstance(value, dict):
        return value
    raise ValueError(f"Expected object or dict, got {type(value).__name__}")


def _require_nonempty_string(data: Dict[str, Any], field_name: str) -> None:
    value = data.get(field_name)
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"FinalReportPayload.{field_name} must be a non-empty string")


def _require_sections_list(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    sections = data.get("sections")
    if not isinstance(sections, list):
        raise ValueError("FinalReportPayload.sections must be a list")
    for section in sections:
        if not isinstance(section, dict):
            raise ValueError("FinalReportPayload.sections must contain objects")
    return sections


def _combined_text(data: Dict[str, Any]) -> str:
    return json.dumps(data, sort_keys=True, ensure_ascii=False).lower()


def _validate_section_references(section: Dict[str, Any]) -> None:
    validate_assembled_report_section(section)
    evidence_ids = set(section["evidence_ids"])
    for ref in section.get("timestamp_refs", []):
        evidence_id = ref.get("evidence_id")
        if evidence_id not in evidence_ids:
            raise ValueError(
                f"FinalReportPayload section {section['section_id']} references "
                f"unknown timestamp evidence_id: {evidence_id}"
            )


def validate_final_report_payload(payload: Any) -> None:
    data = _as_dict(payload)

    for field_name in (
        "report_id",
        "title",
        "report_status",
        "generated_from",
        "evidence_disclaimer",
        "methodology_note",
        "limitations_note",
        "generation_notes",
    ):
        _require_nonempty_string(data, field_name)

    status = data["report_status"]
    if status not in FINAL_REPORT_STATUSES:
        raise ValueError(f"FinalReportPayload.report_status is unsupported: {status}")
    if status == "report_ready" or data.get("report_ready"):
        raise ValueError("Final Report Generation v1 cannot mark evidence as report_ready")

    sections = _require_sections_list(data)
    if status == "review_draft" and not sections:
        raise ValueError(
            "FinalReportPayload.sections must be a non-empty list for review_draft"
        )

    if status == "review_draft":
        review_text = _combined_text(data)
        required_review_phrases = (
            "fixture-based validation artifact",
            "not a public evidentiary conclusion",
            "requires real transcript/timestamp/quote replacement before publication",
            "not for public release",
        )
        for phrase in required_review_phrases:
            if phrase not in review_text:
                raise ValueError(
                    "FinalReportPayload.review_draft must clearly label fixture "
                    f"limitations: missing {phrase!r}"
                )

    if status in {"blocked", "generation_failed"}:
        if "reason" not in data["generation_notes"].lower():
            raise ValueError(
                f"FinalReportPayload.{status} requires a reason in generation_notes"
            )

    for section in sections:
        _validate_section_references(section)

    review_text = _combined_text(data)
    for phrase in FORBIDDEN_REVIEW_TEXT:
        if phrase in review_text:
            raise ValueError(
                f"FinalReportPayload contains forbidden overclaim: {phrase}"
            )


def _coerce_section(data: Dict[str, Any]) -> AssembledReportSection:
    return AssembledReportSection(
        section_id=str(data.get("section_id", "")),
        case_id=str(data.get("case_id", "")),
        claim_ids=list(data.get("claim_ids", [])),
        evidence_ids=list(data.get("evidence_ids", [])),
        heading=str(data.get("heading", "")),
        body=str(data.get("body", "")),
        raw_quotes=list(data.get("raw_quotes", [])),
        timestamp_refs=list(data.get("timestamp_refs", [])),
        assembly_status=str(data.get("assembly_status", "")),
        assembly_notes=str(data.get("assembly_notes", "")),
    )


def load_assembled_report_sections(
    input_path: Path = DEFAULT_REPORT_SECTION_OUTPUT,
) -> List[AssembledReportSection]:
    if not input_path.exists():
        assemble_report_sections_fixture_only(input_path)

    artifact = json.loads(input_path.read_text(encoding="utf-8"))
    sections = artifact.get("sections")
    if not isinstance(sections, list):
        raise ValueError("Report section artifact must contain sections")

    assembled = [_coerce_section(section) for section in sections]
    for section in assembled:
        validate_assembled_report_section(section)
    return assembled


def build_final_report_payload_fixture_only(
    input_path: Path = DEFAULT_REPORT_SECTION_OUTPUT,
) -> FinalReportPayload:
    sections = load_assembled_report_sections(input_path)
    verified_sections = [
        section for section in sections if section.assembly_status == "assembly_verified"
    ]
    if not verified_sections:
        raise ValueError("Final Report Generation v1 requires verified sections")

    payload = FinalReportPayload(
        report_id="DUMA_BOKO_REVIEW_DRAFT_V1_FIXTURE",
        title="Duma Boko Evidence Pipeline Review Draft",
        report_status="review_draft",
        generated_from=str(input_path),
        sections=[section.to_dict() for section in verified_sections],
        evidence_disclaimer=(
            "Review draft fixture-based validation artifact. Not for public release. "
            "Fixture sections are not public evidence and not a public evidentiary "
            "conclusion."
        ),
        methodology_note=(
            "Evidence pipeline demonstration using deterministic fixtures only. "
            "Requires real transcript/timestamp/quote replacement before publication."
        ),
        limitations_note=(
            "This review draft is non-public, non-final, and cannot be used as an "
            "evidentiary conclusion."
        ),
        generation_notes=(
            "Generated for local pipeline validation only. Publication readiness is "
            "not set and release remains gated."
        ),
    )
    validate_final_report_payload(payload)
    return payload


def write_final_report_payload(
    payload: FinalReportPayload,
    output_path: Path = DEFAULT_FINAL_REPORT_PAYLOAD,
) -> Path:
    validate_final_report_payload(payload)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_json(
        {
            "metadata": {
                "generated_at_utc": utc_now_iso(),
                "mode": "fixtures-only",
                "fixture_notice": (
                    "TEST FIXTURE ONLY. Review draft artifact, not public evidence."
                ),
            },
            "payload": payload.to_dict(),
        },
        str(output_path),
    )
    return output_path


def _assert_document_text_allowed(text: str) -> None:
    lower_text = text.lower()
    for phrase in FORBIDDEN_REVIEW_TEXT:
        if phrase in lower_text:
            raise ValueError(f"Generated review document contains forbidden text: {phrase}")


def write_review_docx(
    payload: FinalReportPayload,
    output_path: Path = DEFAULT_FINAL_REPORT_DOCX,
) -> Path:
    validate_final_report_payload(payload)

    document = Document()
    document.add_heading(payload.title, level=0)
    document.add_paragraph("REVIEW DRAFT - NOT FOR PUBLIC RELEASE")
    document.add_paragraph("Fixture-based validation artifact.")
    document.add_paragraph(payload.evidence_disclaimer)
    document.add_heading("Methodology Note", level=1)
    document.add_paragraph(payload.methodology_note)
    document.add_heading("Limitations Note", level=1)
    document.add_paragraph(payload.limitations_note)
    document.add_heading("Case Sections", level=1)

    text_parts = [
        payload.title,
        "REVIEW DRAFT - NOT FOR PUBLIC RELEASE",
        "Fixture-based validation artifact.",
        payload.evidence_disclaimer,
        payload.methodology_note,
        payload.limitations_note,
    ]

    for section in payload.sections:
        document.add_heading(section["heading"], level=2)
        document.add_paragraph(f"Case ID: {section['case_id']}")
        document.add_paragraph(f"Section ID: {section['section_id']}")
        document.add_paragraph(f"Claim IDs: {', '.join(section['claim_ids'])}")
        document.add_paragraph(f"Evidence IDs: {', '.join(section['evidence_ids'])}")
        document.add_paragraph(section["body"])
        document.add_paragraph("Raw quote excerpts:")
        for raw_quote in section["raw_quotes"]:
            document.add_paragraph(raw_quote, style="List Bullet")
        document.add_paragraph("Timestamp references:")
        for ref in section["timestamp_refs"]:
            document.add_paragraph(
                f"{ref['evidence_id']}: {ref['timestamp_start']} to "
                f"{ref['timestamp_end']}",
                style="List Bullet",
            )
        document.add_paragraph(
            "TEST FIXTURE ONLY. Not a public evidentiary conclusion. "
            "Requires real transcript/timestamp/quote replacement before publication."
        )

        text_parts.extend(
            [
                section["heading"],
                section["case_id"],
                section["section_id"],
                " ".join(section["claim_ids"]),
                " ".join(section["evidence_ids"]),
                section["body"],
                " ".join(section["raw_quotes"]),
                " ".join(
                    f"{ref['evidence_id']} {ref['timestamp_start']} {ref['timestamp_end']}"
                    for ref in section["timestamp_refs"]
                ),
            ]
        )

    _assert_document_text_allowed("\n".join(text_parts))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def generate_final_report_fixture_only(
    payload_path: Path = DEFAULT_FINAL_REPORT_PAYLOAD,
    docx_path: Path = DEFAULT_FINAL_REPORT_DOCX,
) -> Dict[str, Any]:
    payload = build_final_report_payload_fixture_only()
    write_final_report_payload(payload, payload_path)
    write_review_docx(payload, docx_path)

    evidence_ids = {
        evidence_id
        for section in payload.sections
        for evidence_id in section["evidence_ids"]
    }
    return {
        "report_id": payload.report_id,
        "report_status": payload.report_status,
        "sections_included": len(payload.sections),
        "evidence_ids_represented": len(evidence_ids),
        "output_docx": str(docx_path),
        "output_payload": str(payload_path),
        "payload": payload,
    }
