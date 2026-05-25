#!/usr/bin/env python3
"""
Validate the active v3.0 Governance Promise-Delivery Divergence Engine.

This test validates the current divergence reconstruction artifact contract.
It does not validate the legacy semantic pipeline.
"""

import json
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document

from analysis.target_search import run_divergence_engine
from evidence.evidence_ingestion import (
    load_or_build_evidence_index,
    load_transcript_artifact,
    upgrade_evidence_status,
)
from evidence.transcript_acquisition import (
    acquire_transcripts,
    extract_youtube_video_id,
    failure_artifact,
    transcript_found_artifact,
    upgrade_evidence_from_transcript_artifact,
    validate_transcript_artifact,
)
from evidence.timestamp_verification import (
    DEFAULT_TIMESTAMP_OUTPUT,
    TimestampCandidate,
    validate_timestamp_candidate,
    verify_timestamps_fixture_only,
)
from evidence.quote_verification import (
    DEFAULT_QUOTE_OUTPUT,
    QuoteCandidate,
    validate_quote_candidate,
    verify_quotes_fixture_only,
)
from evidence.case_evidence_linking import (
    DEFAULT_CASE_LINK_OUTPUT,
    CaseEvidenceLink,
    link_case_evidence_fixture_only,
    validate_case_evidence_link,
    validate_report_sections_resolve,
)
from evidence.report_section_assembly import (
    DEFAULT_REPORT_SECTION_OUTPUT,
    AssembledReportSection,
    assemble_report_sections_fixture_only,
    validate_assembled_report_section,
    validate_assembled_section_resolves,
)
from evidence.final_report_v1 import (
    DEFAULT_FINAL_REPORT_DOCX,
    DEFAULT_FINAL_REPORT_PAYLOAD,
    FinalReportPayload,
    generate_final_report_fixture_only,
    validate_final_report_payload,
)
from evidence.final_report_hardening import (
    DEFAULT_FINAL_REPORT_HARDENING_RECORD,
    DEFAULT_FINAL_REPORT_HARDENING_SUMMARY,
    FinalReportHardeningRecord,
    build_final_report_hardening_record_dry_run,
    harden_final_report_dry_run,
    validate_final_report_hardening_record,
)
from evidence.release_readiness import (
    DEFAULT_RELEASE_READINESS_RECORD,
    DEFAULT_RELEASE_READINESS_SUMMARY,
    DRY_RUN_BLOCKER_REASONS,
    DRY_RUN_RELEASE_NOTES,
    ReleaseReadinessRecord,
    check_release_readiness_dry_run,
    validate_release_readiness_record,
)
from evidence.release_policy import (
    DEFAULT_RELEASE_POLICY_RECORD,
    DEFAULT_RELEASE_POLICY_SUMMARY,
    PROHIBITED_CONDITIONS,
    REQUIRED_CONDITIONS,
    ReleasePolicyRecord,
    check_release_policy_dry_run,
    validate_release_policy_record,
)
from evidence.real_evidence_approval import (
    DEFAULT_REAL_EVIDENCE_APPROVAL_RECORDS_OUTPUT,
    DEFAULT_REAL_EVIDENCE_APPROVAL_SUMMARY_OUTPUT,
    DRY_RUN_BLOCKER_REASONS as APPROVAL_DRY_RUN_BLOCKER_REASONS,
    RealEvidenceApprovalRecord,
    approve_real_evidence_dry_run,
    validate_real_evidence_approval_record,
)
from evidence.real_evidence_inputs import (
    DEFAULT_REAL_EVIDENCE_INPUT_STATUS_OUTPUT,
    DEFAULT_REAL_EVIDENCE_INPUT_SUMMARY_OUTPUT,
    RealEvidenceInputRecord,
    load_real_evidence_input_records,
    validate_real_evidence_input_record,
    validate_real_evidence_inputs_dry_run,
)
from scripts.auto_collect_real_evidence import (
    AUTO_CONTEXT_SUMMARY,
    AUTO_REVIEWER_NOTES,
    AUTO_SPEAKER,
    auto_collect_real_evidence,
    build_candidate_fields,
    collect_candidate_for_record,
    parse_webvtt,
)
from evidence.source_recovery import (
    DEFAULT_SOURCE_RECOVERY_PATH,
    SourceRecoveryCandidateRecord,
    load_source_recovery_candidates,
    summarize_source_recovery_candidates,
    validate_source_recovery_candidate_record,
)
from evidence.recovery_candidate_verification import (
    RecoveryCandidateVerificationRecord,
    validate_recovery_candidate_verification_record,
    verify_candidate_reachability,
    verify_recovery_candidates,
)
from evidence.selected_recovery_source import (
    DEFAULT_SELECTED_RECOVERY_SOURCE_PATH,
    load_selected_recovery_sources,
    summarize_selected_recovery_sources,
    validate_selected_recovery_source_record,
)
from evidence.source_content_extraction import (
    SourceContentExtractionRecord,
    extract_candidate_text,
    extract_content_for_selected_source,
    extract_selected_source_content,
    validate_source_content_extraction_record,
)
from evidence.source_content_verification import (
    build_missing_artifact_records,
    validate_source_content_verification_record,
    verify_extraction_record,
    verify_source_content,
)
from evidence.final_approved_packet import (
    DEFAULT_FINAL_APPROVED_PACKET_RECORD,
    DEFAULT_FINAL_APPROVED_PACKET_SUMMARY,
    DRY_RUN_BLOCKER_REASONS as PACKET_DRY_RUN_BLOCKER_REASONS,
    FinalApprovedEvidencePacketRecord,
    generate_final_approved_packet_dry_run,
    validate_final_approved_packet_record,
)
from evidence.real_evidence_replacement import (
    DEFAULT_REAL_EVIDENCE_OUTPUT_DIR,
    DEFAULT_REAL_QUOTE_OUTPUT,
    DEFAULT_REAL_STATUS_OUTPUT,
    DEFAULT_REAL_TIMESTAMP_OUTPUT,
    RealEvidenceReplacementStatus,
    phrase_matches_real_text,
    replace_real_evidence,
    validate_real_evidence_replacement_status,
)
from evidence.manual_review import (
    DEFAULT_MANUAL_REVIEW_RECORDS_OUTPUT,
    DEFAULT_MANUAL_REVIEW_SUMMARY_OUTPUT,
    DRY_RUN_REVIEW_NOTES,
    ManualReviewRecord,
    manual_review_dry_run,
    validate_manual_review_record,
)
from evidence.evidence_loader import DEFAULT_INDEX_PATH, build_evidence_index, load_seed_evidence
from evidence.evidence_gate import validate_case_evidence_links
from evidence.evidence_schema import (
    CaseObject,
    ClaimObject,
    EvidenceObject,
    ReportSection,
    TranscriptObject,
    validate_case_object,
    validate_claim_object,
    validate_evidence_object,
    validate_transcript_object,
)
from evidence.word_exporter import WordExporter


def placeholder_url_parts():
    token = "example"
    return (
        ".".join((token, "com")),
        token + str(1),
        token + str(2),
        "youtube.com/watch?v=" + token,
    )


def is_placeholder_url(url: str) -> bool:
    return any(part in url.lower() for part in placeholder_url_parts())


def run_command(command: str) -> tuple[int, str]:
    result = subprocess.run(
        command,
        shell=True,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
    )
    print(result.stdout)
    return result.returncode, result.stdout


def assert_nonempty_file(path: str) -> Path:
    p = Path(path)
    if not p.exists():
        raise AssertionError(f"Missing expected file: {path}")
    if p.stat().st_size <= 0:
        raise AssertionError(f"Expected non-empty file: {path}")
    return p


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def text_without_allowed_negated_readiness(text: str) -> str:
    return text.lower().replace("not institution-ready", "")


def assert_value_error(func, expected_text: str) -> None:
    try:
        func()
    except ValueError as exc:
        message = str(exc)
        if expected_text not in message:
            raise AssertionError(
                f"Expected ValueError containing {expected_text!r}, got {message!r}"
            )
        print(f"✓ expected evidence gate failure: {message}")
        return

    raise AssertionError("Expected ValueError was not raised")


def make_reportable_case(url: str, verification_status: str = "source_linked"):
    evidence_id = "CASE_TEST-PROMISE-001"
    evidence = {
        "evidence_id": evidence_id,
        "case_id": "CASE_TEST",
        "source_type": "video",
        "platform": "youtube",
        "title": "Known Source",
        "url": url,
        "evidence_role": "promise",
        "verification_status": verification_status,
        "evidence_strength": "medium",
        "timestamp_start": None,
        "timestamp_end": None,
    }
    return {
        "case_id": "CASE_TEST",
        "evidence_objects": [evidence],
        "claim_evidence_links": {
            "promise": [evidence_id],
            "outcome_or_position": [evidence_id],
            "analysis": [evidence_id],
        },
    }


def validate_divergence_cases(path: Path) -> None:
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    if not isinstance(data, dict):
        raise AssertionError("divergence_cases.json must contain a JSON object")

    if "metadata" not in data:
        raise AssertionError("divergence_cases.json missing root key: metadata")

    cases = data.get("cases")
    if not isinstance(cases, list) or not cases:
        raise AssertionError("divergence_cases.json must contain a non-empty cases list")

    required_case_keys = {
        "case_id",
        "topic",
        "promise",
        "outcome_or_position",
        "divergence_type",
        "analysis",
        "evidence_strength",
        "verification_status",
        "description",
        "raw_urls",
        "evidence_objects",
        "claim_evidence_links",
        "created_at",
        "pipeline_version",
    }

    required_source_keys = {
        "quote",
        "source",
        "url",
        "date",
        "evidence_type",
        "platform",
        "confidence",
    }

    required_evidence_keys = {
        "evidence_id",
        "case_id",
        "source_type",
        "platform",
        "title",
        "url",
        "evidence_role",
        "verification_status",
        "evidence_strength",
    }

    required_claim_links = {
        "promise",
        "outcome_or_position",
        "analysis",
    }

    for idx, case in enumerate(cases):
        if not isinstance(case, dict):
            raise AssertionError(f"Case {idx} must be an object")

        missing = required_case_keys - set(case.keys())
        if missing:
            raise AssertionError(f"Case {idx} missing keys: {sorted(missing)}")

        if case["pipeline_version"] != "divergence_engine_v3.0":
            raise AssertionError(
                f"Case {idx} has wrong pipeline_version: {case['pipeline_version']}"
            )

        for nested_key in ("promise", "outcome_or_position"):
            nested = case[nested_key]
            if not isinstance(nested, dict):
                raise AssertionError(f"Case {idx}.{nested_key} must be an object")

            missing_nested = required_source_keys - set(nested.keys())
            if missing_nested:
                raise AssertionError(
                    f"Case {idx}.{nested_key} missing keys: {sorted(missing_nested)}"
                )

            nested_url = nested.get("url", "")
            if not isinstance(nested_url, str):
                raise AssertionError(f"Case {idx}.{nested_key}.url must be a string")
            if is_placeholder_url(nested_url):
                raise AssertionError(
                    f"Case {idx}.{nested_key}.url must not be a placeholder"
                )

        if not isinstance(case["raw_urls"], list):
            raise AssertionError(f"Case {idx}.raw_urls must be a list")

        for raw_url in case["raw_urls"]:
            if not isinstance(raw_url, str):
                raise AssertionError(f"Case {idx}.raw_urls entries must be strings")
            if is_placeholder_url(raw_url):
                raise AssertionError(f"Case {idx}.raw_urls contains a placeholder")

        evidence_objects = case["evidence_objects"]
        if not isinstance(evidence_objects, list) or not evidence_objects:
            raise AssertionError(f"Case {idx}.evidence_objects must be non-empty")

        evidence_ids = set()
        for evidence_idx, evidence in enumerate(evidence_objects):
            if not isinstance(evidence, dict):
                raise AssertionError(
                    f"Case {idx}.evidence_objects[{evidence_idx}] must be an object"
                )

            missing_evidence = required_evidence_keys - set(evidence.keys())
            if missing_evidence:
                raise AssertionError(
                    f"Case {idx}.evidence_objects[{evidence_idx}] missing keys: "
                    f"{sorted(missing_evidence)}"
                )

            for evidence_key in required_evidence_keys:
                value = evidence[evidence_key]
                if evidence_key == "url":
                    if not isinstance(value, str):
                        raise AssertionError(
                            f"Case {idx}.evidence_objects[{evidence_idx}]."
                            f"{evidence_key} must be a string"
                        )
                    if is_placeholder_url(value):
                        raise AssertionError(
                            f"Case {idx}.evidence_objects[{evidence_idx}]."
                            f"{evidence_key} must not be a placeholder"
                        )
                    continue

                if not isinstance(value, str) or not value.strip():
                    raise AssertionError(
                        f"Case {idx}.evidence_objects[{evidence_idx}]."
                        f"{evidence_key} must be a non-empty string"
                    )

            verification_status = evidence["verification_status"]
            has_timestamp = bool(
                evidence.get("timestamp_start") or evidence.get("timestamp_end")
            )
            if verification_status == "timestamp_verified" and not has_timestamp:
                raise AssertionError(
                    f"Case {idx}.evidence_objects[{evidence_idx}] "
                    "cannot be timestamp_verified without timestamp fields"
                )

            evidence_id = evidence["evidence_id"]
            if evidence_id in evidence_ids:
                raise AssertionError(f"Case {idx} duplicate evidence_id: {evidence_id}")
            evidence_ids.add(evidence_id)

        claim_links = case["claim_evidence_links"]
        if not isinstance(claim_links, dict):
            raise AssertionError(f"Case {idx}.claim_evidence_links must be an object")

        for claim_key in required_claim_links:
            linked_ids = claim_links.get(claim_key)
            if not isinstance(linked_ids, list) or not linked_ids:
                raise AssertionError(
                    f"Case {idx}.claim_evidence_links.{claim_key} must be non-empty"
                )

            for evidence_id in linked_ids:
                if evidence_id not in evidence_ids:
                    raise AssertionError(
                        f"Case {idx}.claim_evidence_links.{claim_key} references "
                        f"unknown evidence_id: {evidence_id}"
                    )

    print(f"✓ divergence_cases.json OK: {len(cases)} cases")


def validate_gate_rejections() -> None:
    assert_value_error(
        lambda: validate_case_evidence_links(make_reportable_case("")),
        "EvidenceObject.url must be a non-empty string",
    )

    token = "example"
    placeholder_url = "https://www.youtube.com/watch?v=" + token + str(1)
    assert_value_error(
        lambda: validate_case_evidence_links(make_reportable_case(placeholder_url)),
        "EvidenceObject.url appears to be a placeholder",
    )

    assert_value_error(
        lambda: validate_case_evidence_links(
            make_reportable_case(
                "https://sources.local/evidence",
                verification_status="timestamp_verified",
            )
        ),
        "requires timestamp_start and timestamp_end",
    )


def validate_schema_objects() -> None:
    evidence = EvidenceObject(
        evidence_id="VID_TEST_001",
        case_id="CASE_TEST",
        source_type="video",
        platform="youtube",
        title="Known Source",
        url="https://www.youtube.com/watch?v=abc123",
        evidence_role="promise_video",
        verification_status="source_found",
        evidence_strength="medium",
    )
    validate_evidence_object(evidence)

    assert_value_error(
        lambda: validate_evidence_object({**evidence.to_dict(), "url": ""}),
        "EvidenceObject.url must be a non-empty string",
    )

    token = "example"
    placeholder_url = "https://www.youtube.com/watch?v=" + token + str(1)
    assert_value_error(
        lambda: validate_evidence_object({**evidence.to_dict(), "url": placeholder_url}),
        "EvidenceObject.url appears to be a placeholder",
    )

    assert_value_error(
        lambda: validate_evidence_object(
            {**evidence.to_dict(), "verification_status": "timestamp_verified"}
        ),
        "requires timestamp_start and timestamp_end",
    )

    assert_value_error(
        lambda: validate_evidence_object(
            {
                **evidence.to_dict(),
                "verification_status": "quote_verified",
                "timestamp_start": "00:01",
                "timestamp_end": "00:10",
                "raw_quote": None,
            }
        ),
        "requires raw_quote",
    )

    claim = ClaimObject(
        claim_id="CLAIM_TEST_001",
        case_id="CASE_TEST",
        claim_type="promise",
        text="A sourced claim.",
        evidence_ids=[evidence.evidence_id],
        verification_status="source_found",
    )
    validate_claim_object(claim)

    assert_value_error(
        lambda: validate_claim_object({**claim.to_dict(), "evidence_ids": []}),
        "ClaimObject.evidence_ids must be a non-empty list",
    )

    case = CaseObject(
        case_id="CASE_TEST",
        title="Test Case",
        domain="governance",
        divergence_type="promise_vs_outcome",
        claims=[claim],
        evidence=[evidence],
        evidence_strength="medium",
        verification_status="source_found",
    )
    validate_case_object(case)

    unresolved_claim = ClaimObject(
        claim_id="CLAIM_TEST_002",
        case_id="CASE_TEST",
        claim_type="promise",
        text="An unresolved claim.",
        evidence_ids=["MISSING_EVIDENCE"],
        verification_status="source_found",
    )
    assert_value_error(
        lambda: validate_case_object({**case.to_dict(), "claims": [unresolved_claim]}),
        "references unknown evidence_id",
    )

    print("✓ EvidenceObject v1 schema validation OK")


def validate_ingestion_lane() -> None:
    index = load_or_build_evidence_index()
    if index["metadata"]["total_evidence_records"] != 2:
        raise AssertionError("Evidence ingestion index must contain 2 records")

    seed_evidence = load_seed_evidence()[0]
    validate_evidence_object(seed_evidence)

    pending_transcript = load_transcript_artifact(seed_evidence.evidence_id)
    if pending_transcript is None:
        raise AssertionError(f"Missing transcript artifact: {seed_evidence.evidence_id}")

    validate_transcript_object(pending_transcript)
    if pending_transcript.transcript_status != "pending":
        raise AssertionError("Seed transcript artifact must remain pending")

    assert_value_error(
        lambda: upgrade_evidence_status(seed_evidence, "timestamp_verified"),
        "Invalid evidence status transition",
    )

    assert_value_error(
        lambda: upgrade_evidence_status(seed_evidence, "transcript_found"),
        "requires a transcript artifact",
    )

    assert_value_error(
        lambda: upgrade_evidence_status(
            seed_evidence,
            "transcript_found",
            pending_transcript,
        ),
        "transcript_status=transcribed",
    )

    transcribed = TranscriptObject(
        evidence_id=seed_evidence.evidence_id,
        transcript_status="transcribed",
        transcript_text="Unit test transcript text.",
        source="unit-test",
        generated_by="unit-test",
        verification_notes="Unit test transcript fixture only.",
    )
    validate_transcript_object(transcribed)

    transcript_found = upgrade_evidence_status(
        seed_evidence,
        "transcript_found",
        transcribed,
    )
    if transcript_found.verification_status != "transcript_found":
        raise AssertionError("Evidence did not move to transcript_found")

    for immutable_field in ("evidence_id", "case_id", "url", "title"):
        if getattr(seed_evidence, immutable_field) != getattr(transcript_found, immutable_field):
            raise AssertionError(f"Evidence ingestion changed {immutable_field}")

    assert_value_error(
        lambda: upgrade_evidence_status(transcript_found, "timestamp_verified"),
        "requires timestamp_start and timestamp_end",
    )

    timestamped = upgrade_evidence_status(
        transcript_found,
        "timestamp_verified",
        timestamp_start="00:01",
        timestamp_end="00:05",
    )
    validate_evidence_object(timestamped)

    assert_value_error(
        lambda: upgrade_evidence_status(timestamped, "quote_verified"),
        "requires raw_quote",
    )

    quoted = upgrade_evidence_status(
        timestamped,
        "quote_verified",
        raw_quote="Unit test quote.",
    )
    validate_evidence_object(quoted)

    report_ready = upgrade_evidence_status(quoted, "report_ready")
    validate_evidence_object(report_ready)

    print("✓ Evidence Ingestion v1 lane validation OK")


def validate_transcript_acquisition_lane() -> None:
    seed_evidence = load_seed_evidence()[0]

    if extract_youtube_video_id(seed_evidence.url) != "e0MLzB5nGDc":
        raise AssertionError("YouTube watch URL video_id extraction failed")
    if extract_youtube_video_id("https://youtu.be/e0MLzB5nGDc") != "e0MLzB5nGDc":
        raise AssertionError("youtu.be URL video_id extraction failed")

    success_artifact = transcript_found_artifact(
        seed_evidence,
        [{"text": "Fixture transcript text for validation only."}],
        "fixture-test",
        "en",
    )
    validate_transcript_artifact(success_artifact)

    assert_value_error(
        lambda: validate_transcript_artifact(
            {**success_artifact.to_dict(), "transcript_text": ""}
        ),
        "transcript_text must be non-empty",
    )

    unavailable_artifact = failure_artifact(
        seed_evidence,
        "transcript_unavailable",
        "fixture-test",
        "No transcript available in fixture.",
    )
    validate_transcript_artifact(unavailable_artifact)

    failed_artifact = failure_artifact(
        seed_evidence,
        "acquisition_failed",
        "fixture-test",
        "Fixture provider failure.",
    )
    validate_transcript_artifact(failed_artifact)

    assert_value_error(
        lambda: validate_transcript_artifact(
            {**unavailable_artifact.to_dict(), "error": ""}
        ),
        "error must be non-empty",
    )

    assert_value_error(
        lambda: validate_transcript_artifact(
            {**failed_artifact.to_dict(), "transcript_text": "not allowed"}
        ),
        "transcript_text must be empty",
    )

    assert_value_error(
        lambda: validate_transcript_artifact(
            {**success_artifact.to_dict(), "verification_status": "timestamp_verified"}
        ),
        "cannot mark evidence as timestamp_verified",
    )

    upgraded = upgrade_evidence_from_transcript_artifact(seed_evidence, success_artifact)
    if upgraded.verification_status != "transcript_found":
        raise AssertionError("Transcript acquisition must only upgrade to transcript_found")
    if upgraded.timestamp_start or upgraded.timestamp_end or upgraded.raw_quote:
        raise AssertionError("Transcript acquisition must not set timestamps or raw_quote")

    assert_value_error(
        lambda: upgrade_evidence_from_transcript_artifact(
            seed_evidence,
            unavailable_artifact,
        ),
        "transcript_found is required",
    )

    summary = acquire_transcripts(fixtures_only=True)
    if summary["processed"] != 2:
        raise AssertionError("Fixture transcript acquisition must process 2 records")
    if summary["transcript_found"] != 0:
        raise AssertionError("Fixture acquisition must not invent transcript text")
    if summary["transcript_unavailable"] != 1 or summary["acquisition_failed"] != 1:
        raise AssertionError("Fixture acquisition counts are not deterministic")

    print("✓ Transcript Acquisition v1 lane validation OK")


def validate_timestamp_verification_lane() -> None:
    candidate = TimestampCandidate(
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        phrase="500 000 new jobs",
        timestamp_start="00:00:10",
        timestamp_end="00:00:18",
        matched_text="We promised to create 500 000 new jobs in five years.",
        match_confidence=1.0,
        verification_status="timestamp_verified",
        verification_notes="Unit test timestamp fixture only.",
    )
    validate_timestamp_candidate(candidate)

    assert_value_error(
        lambda: validate_timestamp_candidate(
            {**candidate.to_dict(), "timestamp_start": ""}
        ),
        "timestamp_start must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_timestamp_candidate(
            {**candidate.to_dict(), "timestamp_end": ""}
        ),
        "timestamp_end must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_timestamp_candidate(
            {
                **candidate.to_dict(),
                "timestamp_start": "00:00:20",
                "timestamp_end": "00:00:10",
            }
        ),
        "timestamp_end must not be before start",
    )

    assert_value_error(
        lambda: validate_timestamp_candidate(
            {**candidate.to_dict(), "matched_text": ""}
        ),
        "matched_text must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_timestamp_candidate(
            {
                **candidate.to_dict(),
                "phrase": "not in transcript",
            }
        ),
        "phrase must appear in matched_text",
    )

    assert_value_error(
        lambda: validate_timestamp_candidate(
            {
                **candidate.to_dict(),
                "evidence_verification_status": "quote_verified",
            }
        ),
        "cannot mark evidence as quote_verified",
    )

    summary = verify_timestamps_fixture_only()
    if summary["processed"] != 2:
        raise AssertionError("Timestamp fixture verification must process 2 records")
    if summary["candidates_found"] != 4:
        raise AssertionError("Timestamp fixture verification must produce 4 candidates")
    if summary["timestamp_verified"] != 4:
        raise AssertionError("Timestamp fixtures must verify 4 timestamp candidates")
    if summary["timestamp_rejected"] != 0 or summary["timestamp_unavailable"] != 0:
        raise AssertionError("Timestamp fixtures should not reject or miss candidates")

    for produced in summary["candidates"]:
        if produced.verification_status != "timestamp_verified":
            raise AssertionError("Fixture candidates must be timestamp_verified only")
        produced_data = produced.to_dict()
        if produced_data.get("raw_quote"):
            raise AssertionError("Timestamp verification must not set raw_quote")
        if produced_data.get("evidence_verification_status") == "quote_verified":
            raise AssertionError("Timestamp verification must not imply quote_verified")

    assert_nonempty_file(str(DEFAULT_TIMESTAMP_OUTPUT))
    print("✓ Timestamp Verification v1 lane validation OK")


def validate_quote_verification_lane() -> None:
    candidate = QuoteCandidate(
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        phrase="500 000 new jobs",
        timestamp_start="00:00:10",
        timestamp_end="00:00:18",
        matched_text="We promised to create 500 000 new jobs in five years.",
        raw_quote="500 000 new jobs",
        quote_confidence=1.0,
        verification_status="quote_verified",
        verification_notes="Unit test quote fixture only. Report readiness pending.",
    )
    validate_quote_candidate(candidate)

    assert_value_error(
        lambda: validate_quote_candidate({**candidate.to_dict(), "raw_quote": ""}),
        "raw_quote must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_quote_candidate(
            {**candidate.to_dict(), "timestamp_start": ""}
        ),
        "timestamp_start must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_quote_candidate(
            {**candidate.to_dict(), "timestamp_end": ""}
        ),
        "timestamp_end must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_quote_candidate({**candidate.to_dict(), "matched_text": ""}),
        "matched_text must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_quote_candidate(
            {**candidate.to_dict(), "raw_quote": "not in transcript"}
        ),
        "raw_quote must be contained in matched_text",
    )

    assert_value_error(
        lambda: validate_quote_candidate(
            {
                **candidate.to_dict(),
                "verification_status": "quote_unavailable",
                "raw_quote": "",
                "verification_notes": "",
            }
        ),
        "verification_notes must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_quote_candidate(
            {
                **candidate.to_dict(),
                "verification_status": "quote_rejected",
                "raw_quote": "",
                "verification_notes": "",
            }
        ),
        "verification_notes must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_quote_candidate(
            {**candidate.to_dict(), "evidence_verification_status": "report_ready"}
        ),
        "cannot mark evidence as report_ready",
    )

    summary = verify_quotes_fixture_only()
    if summary["processed"] != 4:
        raise AssertionError("Quote fixture verification must process 4 timestamps")
    if summary["quote_verified"] != 4:
        raise AssertionError("Quote fixtures must verify 4 quote candidates")
    if (
        summary["quote_candidate_found"] != 0
        or summary["quote_rejected"] != 0
        or summary["quote_unavailable"] != 0
    ):
        raise AssertionError("Quote fixtures should not be pending, rejected, or missed")

    for produced in summary["candidates"]:
        if produced.verification_status != "quote_verified":
            raise AssertionError("Fixture quote candidates must be quote_verified only")
        produced_data = produced.to_dict()
        if produced_data.get("evidence_verification_status") == "report_ready":
            raise AssertionError("Quote verification must not imply report_ready")
        if "TEST FIXTURE ONLY" not in produced.verification_notes:
            raise AssertionError("Fixture quote candidates must be clearly labeled")

    assert_nonempty_file(str(DEFAULT_QUOTE_OUTPUT))
    print("✓ Quote Verification v1 lane validation OK")


def validate_case_evidence_linking_lane() -> None:
    link = CaseEvidenceLink(
        case_id="CASE_002",
        claim_id="CLAIM_JOBS_001",
        evidence_id="VID_JOBS_001",
        quote_id="QUOTE_VID_JOBS_001_500_000_NEW_JOBS",
        phrase="500 000 new jobs",
        timestamp_start="00:00:10",
        timestamp_end="00:00:18",
        raw_quote="500 000 new jobs",
        link_status="link_verified",
        link_notes="Unit test link fixture only. Report readiness pending.",
    )
    validate_case_evidence_link(link)

    for field_name in ("case_id", "claim_id", "evidence_id", "quote_id", "raw_quote"):
        assert_value_error(
            lambda field_name=field_name: validate_case_evidence_link(
                {**link.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    assert_value_error(
        lambda: validate_case_evidence_link(
            {
                **link.to_dict(),
                "link_status": "link_rejected",
                "link_notes": "",
            }
        ),
        "link_notes must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_case_evidence_link(
            {
                **link.to_dict(),
                "link_status": "link_unavailable",
                "link_notes": "",
            }
        ),
        "link_notes must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_case_evidence_link(
            {**link.to_dict(), "evidence_verification_status": "report_ready"}
        ),
        "cannot mark evidence as report_ready",
    )

    seed_evidence = load_seed_evidence()[0]
    unresolved_claim = ClaimObject(
        claim_id="CLAIM_UNRESOLVED",
        case_id=seed_evidence.case_id,
        claim_type="unit_test",
        text="Unit test unresolved evidence claim.",
        evidence_ids=["MISSING_EVIDENCE"],
        verification_status="quote_linked_fixture",
    )
    unresolved_case = CaseObject(
        case_id=seed_evidence.case_id,
        title="Unit test unresolved evidence case",
        domain="unit_test",
        divergence_type="unit_test",
        claims=[unresolved_claim],
        evidence=[seed_evidence],
        evidence_strength=seed_evidence.evidence_strength,
        verification_status="quote_linked_fixture",
    )
    assert_value_error(
        lambda: validate_case_object(unresolved_case),
        "references unknown evidence_id: MISSING_EVIDENCE",
    )

    resolved_claim = ClaimObject(
        claim_id="CLAIM_RESOLVED",
        case_id=seed_evidence.case_id,
        claim_type="unit_test",
        text="Unit test resolved evidence claim.",
        evidence_ids=[seed_evidence.evidence_id],
        verification_status="quote_linked_fixture",
    )
    resolved_case = CaseObject(
        case_id=seed_evidence.case_id,
        title="Unit test resolved evidence case",
        domain="unit_test",
        divergence_type="unit_test",
        claims=[resolved_claim],
        evidence=[seed_evidence],
        evidence_strength=seed_evidence.evidence_strength,
        verification_status="quote_linked_fixture",
    )
    validate_case_object(resolved_case)
    unresolved_section = ReportSection(
        section_id="SECTION_UNRESOLVED",
        case_id=seed_evidence.case_id,
        heading="Unit test unresolved section",
        body="Unit test section body.",
        evidence_ids=["MISSING_EVIDENCE"],
    )
    assert_value_error(
        lambda: validate_report_sections_resolve(
            resolved_case,
            [unresolved_section],
        ),
        "references unknown evidence_id: MISSING_EVIDENCE",
    )

    summary = link_case_evidence_fixture_only()
    if summary["processed"] != 4:
        raise AssertionError("Case linking fixture must process 4 quote candidates")
    if summary["link_verified"] != 2:
        raise AssertionError("Case linking fixtures must verify 2 links")
    if summary["link_unavailable"] != 2:
        raise AssertionError("Case linking fixtures must leave 2 quotes unavailable")
    if summary["link_candidate"] != 0 or summary["link_rejected"] != 0:
        raise AssertionError("Case linking fixtures should not be pending or rejected")
    if summary["cases_built"] != 2:
        raise AssertionError("Case linking fixtures must build 2 cases")
    if summary["claims_built"] != 2:
        raise AssertionError("Case linking fixtures must build 2 claims")
    if summary["report_sections_built"] != 2:
        raise AssertionError("Case linking fixtures must build 2 report sections")

    for produced in summary["links"]:
        produced_data = produced.to_dict()
        if produced_data.get("evidence_verification_status") == "report_ready":
            raise AssertionError("Case evidence linking must not imply report_ready")
        if "TEST FIXTURE ONLY" not in produced.link_notes:
            raise AssertionError("Fixture case links must be clearly labeled")

    assert_nonempty_file(str(DEFAULT_CASE_LINK_OUTPUT))
    print("✓ Case Evidence Linking v1 lane validation OK")


def validate_report_section_assembly_lane() -> None:
    section = AssembledReportSection(
        section_id="SECTION_CASE_002_JOBS",
        case_id="CASE_002",
        claim_ids=["CLAIM_JOBS_001"],
        evidence_ids=["VID_JOBS_001"],
        heading="Jobs Creation Promise Evidence",
        body="TEST FIXTURE ONLY. Unit test assembled section.",
        raw_quotes=["500 000 new jobs"],
        timestamp_refs=[
            {
                "evidence_id": "VID_JOBS_001",
                "timestamp_start": "00:00:10",
                "timestamp_end": "00:00:18",
            }
        ],
        assembly_status="assembly_verified",
        assembly_notes="TEST FIXTURE ONLY. Report readiness pending.",
    )
    validate_assembled_report_section(section)

    for field_name in ("section_id", "case_id", "heading", "body"):
        assert_value_error(
            lambda field_name=field_name: validate_assembled_report_section(
                {**section.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    for field_name in ("claim_ids", "evidence_ids"):
        assert_value_error(
            lambda field_name=field_name: validate_assembled_report_section(
                {**section.to_dict(), field_name: []}
            ),
            f"{field_name} must be a non-empty list",
        )

    assert_value_error(
        lambda: validate_assembled_report_section(
            {**section.to_dict(), "raw_quotes": []}
        ),
        "raw_quotes must be a non-empty list",
    )

    assert_value_error(
        lambda: validate_assembled_report_section(
            {**section.to_dict(), "timestamp_refs": []}
        ),
        "timestamp_refs must be a non-empty list",
    )

    assert_value_error(
        lambda: validate_assembled_report_section(
            {
                **section.to_dict(),
                "assembly_status": "assembly_rejected",
                "raw_quotes": [],
                "timestamp_refs": [],
                "assembly_notes": "",
            }
        ),
        "assembly_notes must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_assembled_report_section(
            {
                **section.to_dict(),
                "assembly_status": "assembly_unavailable",
                "raw_quotes": [],
                "timestamp_refs": [],
                "assembly_notes": "",
            }
        ),
        "assembly_notes must be a non-empty string",
    )

    assert_value_error(
        lambda: validate_assembled_report_section(
            {**section.to_dict(), "evidence_verification_status": "report_ready"}
        ),
        "cannot mark evidence as report_ready",
    )

    seed_evidence = load_seed_evidence()[0]
    claim = ClaimObject(
        claim_id="CLAIM_RESOLVED",
        case_id=seed_evidence.case_id,
        claim_type="unit_test",
        text="Unit test resolved evidence claim.",
        evidence_ids=[seed_evidence.evidence_id],
        verification_status="quote_linked_fixture",
    )
    case = CaseObject(
        case_id=seed_evidence.case_id,
        title="Unit test resolved evidence case",
        domain="unit_test",
        divergence_type="unit_test",
        claims=[claim],
        evidence=[seed_evidence],
        evidence_strength=seed_evidence.evidence_strength,
        verification_status="quote_linked_fixture",
    )
    validate_case_object(case)
    unresolved_section = AssembledReportSection(
        section_id="SECTION_UNRESOLVED",
        case_id=seed_evidence.case_id,
        claim_ids=[claim.claim_id],
        evidence_ids=["MISSING_EVIDENCE"],
        heading="Unit test unresolved assembled section",
        body="TEST FIXTURE ONLY. Unit test assembled section.",
        raw_quotes=["500 000 new jobs"],
        timestamp_refs=[
            {
                "evidence_id": "MISSING_EVIDENCE",
                "timestamp_start": "00:00:10",
                "timestamp_end": "00:00:18",
            }
        ],
        assembly_status="assembly_verified",
        assembly_notes="TEST FIXTURE ONLY. Report readiness pending.",
    )
    assert_value_error(
        lambda: validate_assembled_section_resolves(unresolved_section, case),
        "references unknown evidence_id: MISSING_EVIDENCE",
    )

    summary = assemble_report_sections_fixture_only()
    if summary["processed"] != 4:
        raise AssertionError("Report section assembly must process 4 case links")
    if summary["sections_verified"] != 2:
        raise AssertionError("Report section fixtures must verify 2 sections")
    if summary["sections_unavailable"] != 2:
        raise AssertionError("Report section fixtures must leave 2 sections unavailable")
    if summary["sections_candidate"] != 0 or summary["sections_rejected"] != 0:
        raise AssertionError("Report section fixtures should not be pending or rejected")
    if summary["cases_represented"] != 2:
        raise AssertionError("Report section fixtures must represent 2 cases")
    if summary["evidence_ids_represented"] != 2:
        raise AssertionError("Report section fixtures must represent 2 evidence IDs")

    for produced in summary["sections"]:
        produced_data = produced.to_dict()
        if produced_data.get("evidence_verification_status") == "report_ready":
            raise AssertionError("Report section assembly must not imply report_ready")
        if "TEST FIXTURE ONLY" not in produced.assembly_notes:
            raise AssertionError("Fixture report sections must be clearly labeled")

    assert_nonempty_file(str(DEFAULT_REPORT_SECTION_OUTPUT))
    print("✓ Report Section Assembly v1 lane validation OK")


def validate_final_report_generation_lane() -> None:
    assembly_summary = assemble_report_sections_fixture_only()
    verified_sections = [
        section.to_dict()
        for section in assembly_summary["sections"]
        if section.assembly_status == "assembly_verified"
    ]
    payload = FinalReportPayload(
        report_id="DUMA_BOKO_REVIEW_DRAFT_V1_FIXTURE",
        title="Duma Boko Evidence Pipeline Review Draft",
        report_status="review_draft",
        generated_from=str(DEFAULT_REPORT_SECTION_OUTPUT),
        sections=verified_sections,
        evidence_disclaimer=(
            "Review draft only. Fixture-based validation artifact. Not for public "
            "release. Fixture/test evidence may be present. Fixture sections are "
            "not public evidence and not a public evidentiary conclusion."
        ),
        methodology_note=(
            "Evidence pipeline demonstration using deterministic fixtures only. "
            "Requires real transcript/timestamp/quote replacement before publication. "
            "Evidence conclusions require real transcript, timestamp, quote, and "
            "context approval."
        ),
        limitations_note=(
            "This review draft is non-public, non-final, and cannot be used as an "
            "evidentiary conclusion. Manual review required. Not institution-ready."
        ),
        generation_notes=(
            "Generated for local pipeline validation only. GitHub Actions unavailable "
            "due to account/billing lock; local CI used as validation authority. "
            "Publication readiness is not set and release remains gated."
        ),
    )
    validate_final_report_payload(payload)

    for field_name in (
        "report_id",
        "title",
        "evidence_disclaimer",
        "methodology_note",
        "limitations_note",
        "generation_notes",
    ):
        assert_value_error(
            lambda field_name=field_name: validate_final_report_payload(
                {**payload.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    assert_value_error(
        lambda: validate_final_report_payload({**payload.to_dict(), "sections": []}),
        "sections must be a non-empty list for review_draft",
    )

    assert_value_error(
        lambda: validate_final_report_payload(
            {
                **payload.to_dict(),
                "report_status": "blocked",
                "sections": [],
                "generation_notes": "Blocked pending review.",
            }
        ),
        "blocked requires a reason",
    )

    assert_value_error(
        lambda: validate_final_report_payload(
            {
                **payload.to_dict(),
                "report_status": "generation_failed",
                "sections": [],
                "generation_notes": "Generation failed during review.",
            }
        ),
        "generation_failed requires a reason",
    )

    assert_value_error(
        lambda: validate_final_report_payload({**payload.to_dict(), "report_ready": True}),
        "cannot mark evidence as report_ready",
    )

    summary = generate_final_report_fixture_only()
    if summary["report_status"] != "review_draft":
        raise AssertionError("Final Report Generation v1 must create review_draft only")
    if summary["sections_included"] != 2:
        raise AssertionError("Final Report Generation v1 must include 2 fixture sections")
    if summary["evidence_ids_represented"] != 2:
        raise AssertionError(
            "Final Report Generation v1 must represent 2 fixture evidence IDs"
        )

    payload_path = assert_nonempty_file(str(DEFAULT_FINAL_REPORT_PAYLOAD))
    docx_path = assert_nonempty_file(str(DEFAULT_FINAL_REPORT_DOCX))

    generated_payload = json.loads(payload_path.read_text(encoding="utf-8"))
    generated_text = json.dumps(generated_payload, sort_keys=True).lower()
    docx_text = extract_docx_text(docx_path)
    combined_text = f"{generated_text}\n{docx_text}".lower()

    for required_text in (
        "review draft only",
        "review draft",
        "not for public release",
        "fixture-based validation artifact",
        "fixture/test evidence may be present",
        "manual review required",
        "not institution-ready",
        "not a public evidentiary conclusion",
        "requires real transcript/timestamp/quote replacement before publication",
        "github actions unavailable due to account/billing lock; local ci used as validation authority",
        "evidence conclusions require real transcript, timestamp, quote, and context approval",
    ):
        if required_text not in combined_text:
            raise AssertionError(
                f"Final Report Generation v1 missing required wording: {required_text}"
            )

    forbidden_scan_text = text_without_allowed_negated_readiness(combined_text)
    for forbidden_text in (
        "validated public evidence",
        "final forensic report",
        "institution-ready",
        "public-ready",
        "ready for public release",
        "ready for institutional release",
        "proven corruption",
        "proven failure",
        "report_ready",
    ):
        if forbidden_text in forbidden_scan_text:
            raise AssertionError(
                f"Final Report Generation v1 overclaims with: {forbidden_text}"
            )

    print("✓ Final Report Generation v1 lane validation OK")


def validate_real_evidence_replacement_lane() -> None:
    status = RealEvidenceReplacementStatus(
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        url="https://www.youtube.com/watch?v=e0MLzB5nGDc",
        acquisition_status="manual_review_required",
        transcript_status="unavailable",
        timestamp_status="unavailable",
        quote_status="unavailable",
        manual_review_required=True,
        public_ready=False,
        notes="Manual review required before real evidence replacement.",
    )
    validate_real_evidence_replacement_status(status)

    assert_value_error(
        lambda: validate_real_evidence_replacement_status(
            {**status.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )

    assert_value_error(
        lambda: validate_real_evidence_replacement_status(
            {**status.to_dict(), "manual_review_required": False}
        ),
        "manual_review_required must be true",
    )

    assert_value_error(
        lambda: validate_real_evidence_replacement_status(
            {
                **status.to_dict(),
                "acquisition_status": "transcript_found_real",
                "transcript_status": "transcript_found_real",
            }
        ),
        "transcript_found_real requires a real transcript artifact",
    )

    assert_value_error(
        lambda: validate_real_evidence_replacement_status(
            {**status.to_dict(), "timestamp_status": "timestamp_candidate_real"}
        ),
        "timestamp_candidate_real requires a timestamp candidate",
    )

    assert_value_error(
        lambda: validate_real_evidence_replacement_status(
            {**status.to_dict(), "quote_status": "quote_candidate_real"}
        ),
        "quote_candidate_real requires a quote candidate",
    )

    if not phrase_matches_real_text(
        "500,000 new jobs",
        "We promised to create 500 000 new jobs in five years.",
    ):
        raise AssertionError("Real evidence matching must ignore comma differences")

    empty_transcript_dir = DEFAULT_REAL_EVIDENCE_OUTPUT_DIR / "empty_transcripts"
    summary = replace_real_evidence(
        dry_run=True,
        transcript_dir=empty_transcript_dir,
    )
    if summary["processed"] != 2:
        raise AssertionError("Real Evidence Replacement v1 dry-run must process 2 records")
    if summary["transcript_found_real"] != 0:
        raise AssertionError("Dry-run without real transcript artifacts must not find transcripts")
    if summary["timestamp_candidate_real"] != 0:
        raise AssertionError("Dry-run without real transcripts must not create timestamps")
    if summary["quote_candidate_real"] != 0:
        raise AssertionError("Dry-run without real transcripts must not create quotes")
    if summary["manual_review_required"] != 2:
        raise AssertionError("Dry-run must keep both seed records in manual review")

    assert_nonempty_file(str(DEFAULT_REAL_STATUS_OUTPUT))
    assert_nonempty_file(str(DEFAULT_REAL_TIMESTAMP_OUTPUT))
    assert_nonempty_file(str(DEFAULT_REAL_QUOTE_OUTPUT))
    if empty_transcript_dir.exists() and any(empty_transcript_dir.iterdir()):
        raise AssertionError("Dry-run must not write real transcript artifacts")

    print("✓ Real Evidence Replacement v1 lane validation OK")


def validate_manual_review_lane() -> None:
    record = ManualReviewRecord(
        review_id="REVIEW_VID_JOBS_001",
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        url="https://www.youtube.com/watch?v=e0MLzB5nGDc",
        reviewer="manual-review-v1-test",
        review_status="pending_review",
        transcript_review="Transcript requires human verification.",
        timestamp_review="Timestamp requires human verification.",
        quote_review="Quote requires human verification.",
        context_review="Context requires human verification.",
        speaker_review="Speaker confidence requires human verification.",
        publication_recommendation="needs_more_evidence",
        manual_review_required=True,
        public_ready=False,
        report_ready=False,
        review_notes=DRY_RUN_REVIEW_NOTES,
    )
    validate_manual_review_record(record)

    for field_name in ("review_id", "evidence_id", "reviewer", "review_notes"):
        assert_value_error(
            lambda field_name=field_name: validate_manual_review_record(
                {**record.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    assert_value_error(
        lambda: validate_manual_review_record(
            {**record.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )

    assert_value_error(
        lambda: validate_manual_review_record(
            {**record.to_dict(), "report_ready": True}
        ),
        "report_ready must be false",
    )

    assert_value_error(
        lambda: validate_manual_review_record(
            {**record.to_dict(), "manual_review_required": False}
        ),
        "manual_review_required must be true",
    )

    assert_value_error(
        lambda: validate_manual_review_record(
            {
                **record.to_dict(),
                "review_status": "reviewed_rejected",
                "publication_recommendation": "do_not_publish",
                "review_notes": "Rejected by dry-run reviewer.",
            }
        ),
        "reviewed_rejected requires a reason",
    )

    assert_value_error(
        lambda: validate_manual_review_record(
            {
                **record.to_dict(),
                "review_status": "reviewed_with_concerns",
                "review_notes": "Concern noted by dry-run reviewer.",
            }
        ),
        "reviewed_with_concerns requires a reason",
    )

    hardening_candidate = ManualReviewRecord(
        **{
            **record.to_dict(),
            "review_status": "reviewed_candidate",
            "publication_recommendation": "candidate_for_hardening",
            "review_notes": (
                "Dry-run candidate for hardening only. Manual review still required."
            ),
        }
    )
    validate_manual_review_record(hardening_candidate)
    if hardening_candidate.public_ready or hardening_candidate.report_ready:
        raise AssertionError("candidate_for_hardening must not imply readiness")

    summary = manual_review_dry_run()
    if summary["total_evidence_records_reviewed"] != 2:
        raise AssertionError("Manual Review v1 dry-run must produce 2 records")
    if summary["pending_review"] != 2:
        raise AssertionError("Manual Review v1 dry-run records must remain pending")
    if summary["public_ready"] != 0 or summary["report_ready"] != 0:
        raise AssertionError("Manual Review v1 dry-run must not mark readiness")
    for produced in summary["records"]:
        if produced.manual_review_required is not True:
            raise AssertionError("Manual Review v1 must require manual review")
        if produced.public_ready or produced.report_ready:
            raise AssertionError("Manual Review v1 must not mark readiness")
        if produced.publication_recommendation != "needs_more_evidence":
            raise AssertionError("Manual Review v1 dry-run must need more evidence")

    assert_nonempty_file(str(DEFAULT_MANUAL_REVIEW_RECORDS_OUTPUT))
    assert_nonempty_file(str(DEFAULT_MANUAL_REVIEW_SUMMARY_OUTPUT))
    print("✓ Manual Review v1 lane validation OK")


def validate_final_report_hardening_lane() -> None:
    record = FinalReportHardeningRecord(
        hardening_id="FINAL_REPORT_HARDENING_TEST",
        report_id="DUMA_BOKO_REVIEW_DRAFT_V1_FIXTURE",
        hardening_status="hardened_review_draft",
        release_status="manual_review_required",
        methodology_note=(
            "Review draft only. Hardening adds methodology and audit context "
            "without changing readiness."
        ),
        limitations_note=(
            "Not for public release. Fixture/test evidence may be present. "
            "Manual review required. Not institution-ready."
        ),
        evidence_status_summary={
            "sections_included": 2,
            "evidence_ids_represented": 2,
        },
        manual_review_summary={
            "pending_review": 2,
            "public_ready": 0,
            "report_ready": 0,
        },
        source_index=[
            {
                "section_id": "SECTION_CASE_002_JOBS",
                "evidence_ids": ["VID_JOBS_001"],
            }
        ],
        audit_trail_summary=["Local CI validated the dry-run lane."],
        ci_validation_note="Local CI used as validation authority.",
        github_actions_note=(
            "GitHub Actions unavailable due to account/billing lock; local CI used "
            "as validation authority."
        ),
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        hardening_notes=(
            "Dry-run hardening record only. Review draft remains blocked from "
            "release pending real-source evidence and manual approval."
        ),
    )
    validate_final_report_hardening_record(record)

    for field_name in (
        "hardening_id",
        "report_id",
        "methodology_note",
        "limitations_note",
        "ci_validation_note",
        "github_actions_note",
        "hardening_notes",
    ):
        assert_value_error(
            lambda field_name=field_name: validate_final_report_hardening_record(
                {**record.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    for field_name in ("evidence_status_summary", "manual_review_summary"):
        assert_value_error(
            lambda field_name=field_name: validate_final_report_hardening_record(
                {**record.to_dict(), field_name: {}}
            ),
            f"{field_name} must be a non-empty object",
        )

    for field_name in ("source_index", "audit_trail_summary"):
        assert_value_error(
            lambda field_name=field_name: validate_final_report_hardening_record(
                {**record.to_dict(), field_name: []}
            ),
            f"{field_name} must be a non-empty list",
        )

    assert_value_error(
        lambda: validate_final_report_hardening_record(
            {**record.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )

    assert_value_error(
        lambda: validate_final_report_hardening_record(
            {**record.to_dict(), "institutional_ready": True}
        ),
        "institutional_ready must be false",
    )

    assert_value_error(
        lambda: validate_final_report_hardening_record(
            {**record.to_dict(), "report_ready": True}
        ),
        "report_ready must be false",
    )

    assert_value_error(
        lambda: validate_final_report_hardening_record(
            {
                **record.to_dict(),
                "hardening_status": "hardening_blocked",
                "release_status": "blocked_from_public_release",
                "hardening_notes": "Blocked pending manual review.",
            }
        ),
        "hardening_blocked requires a reason",
    )

    assert_value_error(
        lambda: validate_final_report_hardening_record(
            {
                **record.to_dict(),
                "hardening_notes": "This draft is ready for institutional release.",
            }
        ),
        "prohibited readiness claim",
    )

    hardened = FinalReportHardeningRecord(**record.to_dict())
    validate_final_report_hardening_record(hardened)
    if hardened.public_ready or hardened.institutional_ready or hardened.report_ready:
        raise AssertionError("hardened_review_draft must not imply readiness")

    built = build_final_report_hardening_record_dry_run()
    validate_final_report_hardening_record(built)
    summary = harden_final_report_dry_run()
    if summary["hardening_status"] != "hardened_review_draft":
        raise AssertionError("Final Report Hardening v1 must harden review draft only")
    if summary["release_status"] != "manual_review_required":
        raise AssertionError("Final Report Hardening v1 must require manual review")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Final Report Hardening v1 must not mark readiness")

    assert_nonempty_file(str(DEFAULT_FINAL_REPORT_HARDENING_RECORD))
    assert_nonempty_file(str(DEFAULT_FINAL_REPORT_HARDENING_SUMMARY))
    print("✓ Final Report Hardening v1 lane validation OK")


def validate_release_readiness_lane() -> None:
    record = ReleaseReadinessRecord(
        readiness_id="RELEASE_READINESS_TEST",
        report_id="DUMA_BOKO_REVIEW_DRAFT_V1_FIXTURE",
        release_status="blocked_manual_review_required",
        release_blocked=True,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        evidence_approval_status="manual_review_required",
        manual_review_status="pending_review",
        source_verification_status="source_present_but_not_release_approved",
        timestamp_verification_status="fixture_verified_not_public_evidence",
        quote_verification_status="fixture_verified_not_public_evidence",
        context_verification_status="manual_context_review_required",
        ci_policy_status="github_actions_unavailable_local_ci_used",
        blocker_reasons=list(DRY_RUN_BLOCKER_REASONS),
        release_notes=DRY_RUN_RELEASE_NOTES,
    )
    validate_release_readiness_record(record)

    for field_name in (
        "readiness_id",
        "report_id",
        "release_status",
        "release_notes",
        "evidence_approval_status",
        "manual_review_status",
        "source_verification_status",
        "timestamp_verification_status",
        "quote_verification_status",
        "context_verification_status",
        "ci_policy_status",
    ):
        assert_value_error(
            lambda field_name=field_name: validate_release_readiness_record(
                {**record.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    assert_value_error(
        lambda: validate_release_readiness_record(
            {**record.to_dict(), "blocker_reasons": []}
        ),
        "blocker_reasons must be non-empty",
    )

    assert_value_error(
        lambda: validate_release_readiness_record(
            {**record.to_dict(), "release_blocked": False}
        ),
        "release_blocked must be true",
    )

    assert_value_error(
        lambda: validate_release_readiness_record(
            {**record.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )

    assert_value_error(
        lambda: validate_release_readiness_record(
            {**record.to_dict(), "institutional_ready": True}
        ),
        "institutional_ready must be false",
    )

    assert_value_error(
        lambda: validate_release_readiness_record(
            {**record.to_dict(), "report_ready": True}
        ),
        "report_ready must be false",
    )

    assert_value_error(
        lambda: validate_release_readiness_record(
            {
                **record.to_dict(),
                "release_notes": "This report is ready for public release.",
            }
        ),
        "prohibited release claim",
    )

    pending = ReleaseReadinessRecord(
        **{
            **record.to_dict(),
            "release_status": "release_candidate_pending_approval",
            "blocker_reasons": [
                "release candidate remains blocked pending explicit future approvals"
            ],
        }
    )
    validate_release_readiness_record(pending)
    if pending.public_ready or pending.institutional_ready or pending.report_ready:
        raise AssertionError(
            "release_candidate_pending_approval must not imply readiness"
        )

    summary = check_release_readiness_dry_run()
    if summary["release_status"] != "blocked_manual_review_required":
        raise AssertionError("Release Readiness v1 must block manual-review drafts")
    if summary["release_blocked"] is not True:
        raise AssertionError("Release Readiness v1 must block release")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Release Readiness v1 must not mark readiness")
    if summary["blocker_count"] != len(DRY_RUN_BLOCKER_REASONS):
        raise AssertionError("Release Readiness v1 blocker count changed")
    if summary["record"].blocker_reasons != DRY_RUN_BLOCKER_REASONS:
        raise AssertionError("Release Readiness v1 blockers must be deterministic")

    assert_nonempty_file(str(DEFAULT_RELEASE_READINESS_RECORD))
    assert_nonempty_file(str(DEFAULT_RELEASE_READINESS_SUMMARY))
    print("✓ Release Readiness v1 lane validation OK")


def validate_release_policy_lane() -> None:
    record = ReleasePolicyRecord(
        policy_id="RELEASE_POLICY_TEST",
        policy_status="active_local_ci_override_policy",
        local_ci_authority="temporary CI validation replacement only",
        github_actions_status="unavailable_due_to_billing_account_lock",
        manual_override_allowed=True,
        required_conditions=list(REQUIRED_CONDITIONS),
        prohibited_conditions=list(PROHIBITED_CONDITIONS),
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        policy_notes=(
            "Manual override is limited to CI validation replacement and is not "
            "evidence/public/institutional release approval."
        ),
    )
    validate_release_policy_record(record)

    for field_name in (
        "policy_id",
        "policy_status",
        "local_ci_authority",
        "github_actions_status",
        "policy_notes",
    ):
        assert_value_error(
            lambda field_name=field_name: validate_release_policy_record(
                {**record.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    for field_name in ("required_conditions", "prohibited_conditions"):
        assert_value_error(
            lambda field_name=field_name: validate_release_policy_record(
                {**record.to_dict(), field_name: []}
            ),
            f"{field_name} must be a non-empty list",
        )

    assert_value_error(
        lambda: validate_release_policy_record(
            {**record.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )

    assert_value_error(
        lambda: validate_release_policy_record(
            {**record.to_dict(), "institutional_ready": True}
        ),
        "institutional_ready must be false",
    )

    assert_value_error(
        lambda: validate_release_policy_record(
            {**record.to_dict(), "report_ready": True}
        ),
        "report_ready must be false",
    )

    assert_value_error(
        lambda: validate_release_policy_record(
            {
                **record.to_dict(),
                "policy_notes": "Manual override approves evidence release.",
            }
        ),
        "must not imply release approval",
    )

    validate_release_policy_record(record)
    if record.public_ready or record.institutional_ready or record.report_ready:
        raise AssertionError("Release Policy v1 must not imply readiness")
    if "CI validation replacement" not in record.local_ci_authority:
        raise AssertionError("Release Policy v1 must be limited to CI validation")

    summary = check_release_policy_dry_run()
    if summary["manual_override_allowed"] is not True:
        raise AssertionError("Release Policy v1 dry-run must allow CI override only")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Release Policy v1 dry-run must not mark readiness")
    if "billing_account_lock" not in summary["github_actions_status"]:
        raise AssertionError("Release Policy v1 must document GitHub Actions lock")

    readiness_summary = check_release_readiness_dry_run()
    if readiness_summary["release_blocked"] is not True:
        raise AssertionError("Release Readiness v1 must remain blocked")

    assert_nonempty_file(str(DEFAULT_RELEASE_POLICY_RECORD))
    assert_nonempty_file(str(DEFAULT_RELEASE_POLICY_SUMMARY))
    print("✓ Release Policy v1 lane validation OK")


def validate_real_evidence_inputs_lane() -> None:
    pending = RealEvidenceInputRecord(
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        source_url="https://www.youtube.com/watch?v=e0MLzB5nGDc",
        transcript_text="",
        timestamp_start="",
        timestamp_end="",
        quote_text="",
        speaker="",
        context_summary="",
        case_relevance_note="",
        reviewer="",
        reviewer_notes="",
        verification_status="pending_human_entry",
    )
    validate_real_evidence_input_record(pending)

    verified = RealEvidenceInputRecord(
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        source_url="https://www.youtube.com/watch?v=e0MLzB5nGDc",
        transcript_text="Human-entered transcript excerpt.",
        timestamp_start="00:01:00",
        timestamp_end="00:01:10",
        quote_text="Human-entered quote excerpt.",
        speaker="Human-verified speaker",
        context_summary="Human-entered context summary.",
        case_relevance_note="Human-entered case relevance note.",
        reviewer="real-evidence-population-v1-test",
        reviewer_notes="Human reviewer notes for approval review.",
        verification_status="verified_for_approval_review",
    )
    validate_real_evidence_input_record(verified)

    for field_name in (
        "transcript_text",
        "timestamp_start",
        "timestamp_end",
        "quote_text",
        "speaker",
        "reviewer",
    ):
        assert_value_error(
            lambda field_name=field_name: validate_real_evidence_input_record(
                {**verified.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    assert_value_error(
        lambda: validate_real_evidence_input_record(
            {**verified.to_dict(), "verification_status": "unsupported_status"}
        ),
        "verification_status is unsupported",
    )

    templates = load_real_evidence_input_records()
    if len(templates) != 2:
        raise AssertionError("Real Evidence Population v1 must provide 2 templates")
    for record in templates:
        if record.verification_status != "pending_human_entry":
            raise AssertionError("Real Evidence Population v1 templates must be pending")
        for field_name in (
            "transcript_text",
            "timestamp_start",
            "timestamp_end",
            "quote_text",
            "speaker",
            "context_summary",
            "case_relevance_note",
            "reviewer",
            "reviewer_notes",
        ):
            if getattr(record, field_name) != "":
                raise AssertionError(
                    "Real Evidence Population v1 templates must keep human "
                    f"entry field empty: {field_name}"
                )

    summary = validate_real_evidence_inputs_dry_run()
    if summary["total_input_records"] != 2:
        raise AssertionError("Real Evidence Population v1 dry-run must validate 2 records")
    if summary["pending_human_entry"] != 2:
        raise AssertionError("Real Evidence Population v1 templates must remain pending")
    if summary["records_ready_for_approval_review"] != 0:
        raise AssertionError("Real Evidence Population v1 must not auto-approve inputs")

    approval_summary = approve_real_evidence_dry_run()
    if approval_summary["approved_evidence_candidate"] != 0:
        raise AssertionError("Real Evidence Approval v1 must still block fixtures")
    if approval_summary["blocked_manual_review_required"] != 2:
        raise AssertionError("Real Evidence Approval v1 must still block both records")

    packet_summary = generate_final_approved_packet_dry_run()
    if packet_summary["packet_status"] != "blocked_no_approved_evidence":
        raise AssertionError("Final Approved Evidence Packet v1 must remain blocked")
    if packet_summary["approved_evidence_count"] != 0:
        raise AssertionError("Final Approved Evidence Packet v1 must not approve")

    assert_nonempty_file(str(DEFAULT_REAL_EVIDENCE_INPUT_STATUS_OUTPUT))
    assert_nonempty_file(str(DEFAULT_REAL_EVIDENCE_INPUT_SUMMARY_OUTPUT))
    print("✓ Real Evidence Population v1 input validation OK")


def validate_real_evidence_auto_collection_helper_lane() -> None:
    vtt_text = """WEBVTT

00:00:01.000 --> 00:00:03.500
Hello <c>world</c>

00:00:04.000 --> 00:00:08.000
This is the longer caption line for candidate selection.
"""
    segments = parse_webvtt(vtt_text)
    if len(segments) != 2:
        raise AssertionError("Auto-collection VTT parser must parse 2 segments")
    if segments[0].timestamp_start != "00:01":
        raise AssertionError("Auto-collection parser must format start timestamps")
    if segments[1].timestamp_end != "00:08":
        raise AssertionError("Auto-collection parser must format end timestamps")
    if segments[0].text != "Hello world":
        raise AssertionError("Auto-collection parser must clean caption markup")

    record = {
        "evidence_id": "VID_JOBS_001",
        "case_id": "CASE_002",
        "source_url": "https://www.youtube.com/watch?v=e0MLzB5nGDc",
        "transcript_text": "",
        "timestamp_start": "",
        "timestamp_end": "",
        "quote_text": "",
        "speaker": "",
        "context_summary": "",
        "case_relevance_note": "",
        "reviewer": "",
        "reviewer_notes": "",
        "verification_status": "pending_human_entry",
    }
    candidate_fields = build_candidate_fields(record, segments)
    if candidate_fields["verification_status"] != "entered_pending_review":
        raise AssertionError("Auto-collection candidate must remain pending review")
    if candidate_fields["speaker"] != AUTO_SPEAKER:
        raise AssertionError("Auto-collection must use an unverified speaker marker")
    if candidate_fields["context_summary"] != AUTO_CONTEXT_SUMMARY:
        raise AssertionError("Auto-collection must require human context review")
    if candidate_fields["reviewer"] != "":
        raise AssertionError("Auto-collection must not invent a reviewer")
    if candidate_fields["reviewer_notes"] != AUTO_REVIEWER_NOTES:
        raise AssertionError("Auto-collection must require human verification")
    if "public_ready" in candidate_fields or "report_ready" in candidate_fields:
        raise AssertionError("Auto-collection must not emit readiness fields")

    no_network_status = collect_candidate_for_record(
        record, no_network=True, timeout_seconds=1
    )
    if no_network_status.collection_status != "network_disabled":
        raise AssertionError("Auto-collection --no-network must report network_disabled")
    if no_network_status.candidate_available:
        raise AssertionError("Auto-collection --no-network must not create candidates")
    if (
        no_network_status.public_ready
        or no_network_status.institutional_ready
        or no_network_status.report_ready
    ):
        raise AssertionError("Auto-collection status must not mark readiness")

    missing_tool_status = collect_candidate_for_record(
        record, no_network=False, timeout_seconds=1, yt_dlp_path=""
    )
    if missing_tool_status.collection_status != "yt_dlp_missing":
        raise AssertionError("Auto-collection missing yt-dlp must be non-fatal")

    template_path = Path("data/real_evidence_inputs/VID_JOBS_001.template.json")
    before = template_path.read_text(encoding="utf-8")
    summary = auto_collect_real_evidence(
        evidence_id="VID_JOBS_001",
        dry_run=True,
        no_network=True,
        timeout_seconds=1,
    )
    after = template_path.read_text(encoding="utf-8")
    if before != after:
        raise AssertionError("Auto-collection dry-run must not modify templates")
    if summary["selected_evidence_count"] != 1:
        raise AssertionError("Auto-collection evidence-id filter must select one record")
    if summary["candidate_count"] != 0 or summary["unavailable_count"] != 1:
        raise AssertionError("Auto-collection --no-network summary changed")
    if summary["write_count"] != 0 or summary["backup_count"] != 0:
        raise AssertionError("Auto-collection dry-run must not write or backup")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Auto-collection summary must not mark readiness")

    print("✓ Real Evidence Auto-Collection Helper v1 validation OK")


def validate_source_recovery_lane() -> None:
    records = load_source_recovery_candidates()
    if len(records) != 5:
        raise AssertionError("Source Recovery v1 must load 5 candidate records")

    original_ids = {record.original_evidence_id for record in records}
    if original_ids != {"VID_JOBS_001", "VID_HEALTH_001"}:
        raise AssertionError("Source Recovery v1 must preserve original evidence IDs")

    summary = summarize_source_recovery_candidates(records)
    if summary["total_candidates"] != 5:
        raise AssertionError("Source Recovery v1 candidate count changed")
    if summary["jobs_candidates"] != 2:
        raise AssertionError("Source Recovery v1 jobs candidate count changed")
    if summary["health_candidates"] != 3:
        raise AssertionError("Source Recovery v1 health candidate count changed")
    if summary["approved_candidates"] != 0:
        raise AssertionError("Source Recovery v1 must not approve candidates")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Source Recovery v1 must not mark readiness")

    for record in records:
        data = record.to_dict()
        validate_source_recovery_candidate_record(record)
        if record.verification_status != "candidate_unverified":
            raise AssertionError("Seed source recovery candidates must be unverified")
        for forbidden_field in (
            "approved",
            "approved_candidate",
            "approval_candidate",
            "approved_evidence_candidate",
            "public_ready",
            "institutional_ready",
            "report_ready",
        ):
            if forbidden_field in data:
                raise AssertionError(
                    "Source Recovery v1 must not include approval/readiness fields"
                )

    candidate = records[0]
    assert_value_error(
        lambda: validate_source_recovery_candidate_record(
            {**candidate.to_dict(), "verification_status": "approved_evidence"}
        ),
        "verification_status is unsupported",
    )
    assert_value_error(
        lambda: validate_source_recovery_candidate_record(
            {**candidate.to_dict(), "recovery_source_url": ""}
        ),
        "recovery_source_url must be a non-empty string",
    )
    assert_value_error(
        lambda: validate_source_recovery_candidate_record(
            {**candidate.to_dict(), "public_ready": False}
        ),
        "must not contain approval/readiness field",
    )
    assert_value_error(
        lambda: validate_source_recovery_candidate_record(
            {
                **candidate.to_dict(),
                "original_evidence_id": "VID_UNKNOWN_001",
            }
        ),
        "original_evidence_id is unsupported",
    )

    with tempfile.TemporaryDirectory() as temp_dir:
        payload = json.loads(DEFAULT_SOURCE_RECOVERY_PATH.read_text(encoding="utf-8"))
        duplicate = json.loads(json.dumps(payload))
        duplicate["candidates"].append(dict(duplicate["candidates"][0]))
        duplicate_path = Path(temp_dir) / "source_recovery_candidates.json"
        duplicate_path.write_text(json.dumps(duplicate, indent=2), encoding="utf-8")
        assert_value_error(
            lambda: load_source_recovery_candidates(path=duplicate_path),
            "Duplicate source recovery candidate ID",
        )

    for command in (
        "python scripts/inspect_source_recovery.py",
        "python scripts/inspect_source_recovery.py --evidence-id VID_JOBS_001",
        "python scripts/inspect_source_recovery.py --evidence-id VID_HEALTH_001",
    ):
        exit_code, output = run_command(command)
        if exit_code != 0:
            raise AssertionError(f"Source recovery inspection command failed: {command}")
        if "approved_candidates: 0" not in output:
            raise AssertionError("Source recovery inspection must not approve candidates")
        if "public_ready: False" not in output:
            raise AssertionError("Source recovery inspection must not mark public readiness")

    print("✓ Evidence Source Recovery v1 lane validation OK")


def validate_recovery_candidate_verification_lane() -> None:
    candidates = load_source_recovery_candidates()
    if len(candidates) != 5:
        raise AssertionError("Recovery Candidate Verification v1 must load 5 candidates")

    no_network_record = verify_candidate_reachability(
        candidates[0],
        no_network=True,
        timeout_seconds=1,
    )
    validate_recovery_candidate_verification_record(no_network_record)
    if no_network_record.reachability_status != "not_checked":
        raise AssertionError("No-network verification must not check reachability")
    if no_network_record.content_status != "requires_manual_review":
        raise AssertionError("No-network verification must require manual review")
    if no_network_record.verification_status != "candidate_unverified":
        raise AssertionError("No-network verification must remain unverified")
    if (
        no_network_record.verified_for_manual_review
        or no_network_record.approved_evidence
        or no_network_record.public_ready
        or no_network_record.institutional_ready
        or no_network_record.report_ready
    ):
        raise AssertionError("Recovery candidate verification must not mark readiness")

    assert_value_error(
        lambda: validate_recovery_candidate_verification_record(
            {**no_network_record.to_dict(), "approved_evidence": True}
        ),
        "approved_evidence must be false",
    )
    assert_value_error(
        lambda: validate_recovery_candidate_verification_record(
            {**no_network_record.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )
    assert_value_error(
        lambda: validate_recovery_candidate_verification_record(
            {**no_network_record.to_dict(), "recovery_source_url": ""}
        ),
        "recovery_source_url must be a non-empty string",
    )

    protected_paths = [
        Path("data/evidence_seed/videos.yaml"),
        Path("data/real_evidence_inputs/VID_JOBS_001.template.json"),
        Path("data/real_evidence_inputs/VID_HEALTH_001.template.json"),
    ]
    before = {
        path: path.read_text(encoding="utf-8")
        for path in protected_paths
    }
    summary = verify_recovery_candidates(no_network=True, timeout_seconds=1)
    after = {
        path: path.read_text(encoding="utf-8")
        for path in protected_paths
    }
    if before != after:
        raise AssertionError(
            "Recovery candidate verification must not modify protected source files"
        )
    if summary["total_candidates"] != 5:
        raise AssertionError("Recovery verification no-network must process 5 records")
    if summary["not_checked_candidates"] != 5:
        raise AssertionError("Recovery verification no-network must not check URLs")
    if summary["requires_manual_review"] != 5:
        raise AssertionError("Recovery verification no-network must require review")
    if summary["approved_evidence"] != 0 or summary["verified_for_manual_review"] != 0:
        raise AssertionError("Recovery verification must not approve candidates")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Recovery verification must not mark readiness")

    filtered_summary = verify_recovery_candidates(
        candidate_id="RECOVERY_VID_JOBS_001_UDC_HOME",
        no_network=True,
        timeout_seconds=1,
    )
    if filtered_summary["total_candidates"] != 1:
        raise AssertionError("Recovery verification candidate-id filter failed")

    exit_code, output = run_command("python scripts/verify_recovery_candidates.py --no-network")
    if exit_code != 0:
        raise AssertionError("Recovery verification no-network CLI failed")
    if "approved_evidence: 0" not in output:
        raise AssertionError("Recovery verification CLI must not approve evidence")
    if "public_ready: False" not in output:
        raise AssertionError("Recovery verification CLI must not mark readiness")

    print("✓ Recovery Candidate Verification v1 lane validation OK")


def validate_selected_recovery_source_lane() -> None:
    candidates = load_source_recovery_candidates()
    records = load_selected_recovery_sources(candidates=candidates)
    if len(records) != 2:
        raise AssertionError("Selected Recovery Source v1 must load 2 records")

    candidate_ids = {candidate.recovery_candidate_id for candidate in candidates}
    selected_by_evidence_id = {
        record.original_evidence_id: record
        for record in records
        if record.selection_status == "selected_for_content_review"
    }
    if set(selected_by_evidence_id) != {"VID_JOBS_001", "VID_HEALTH_001"}:
        raise AssertionError("Selected Recovery Source v1 must select both evidence IDs")

    jobs_record = selected_by_evidence_id["VID_JOBS_001"]
    health_record = selected_by_evidence_id["VID_HEALTH_001"]
    if jobs_record.selected_recovery_candidate_id != "RECOVERY_VID_JOBS_001_UDC_HOME":
        raise AssertionError("Selected jobs recovery source changed")
    if health_record.selected_recovery_candidate_id != "RECOVERY_VID_HEALTH_001_REUTERS":
        raise AssertionError("Selected health recovery source changed")

    summary = summarize_selected_recovery_sources(records)
    if summary["selected_count"] != 2:
        raise AssertionError("Selected Recovery Source v1 selected count changed")
    if summary["pending_selection_count"] != 0:
        raise AssertionError("Selected Recovery Source v1 should have no pending rows")
    if summary["approved_evidence"] != 0:
        raise AssertionError("Selected Recovery Source v1 must not approve evidence")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Selected Recovery Source v1 must not mark readiness")

    for record in records:
        validate_selected_recovery_source_record(record, candidates)
        if record.selected_recovery_candidate_id not in candidate_ids:
            raise AssertionError("Selected source candidate ID must exist in registry")
        if (
            record.approved_evidence
            or record.public_ready
            or record.institutional_ready
            or record.report_ready
        ):
            raise AssertionError("Selected source must not approve or mark readiness")

    assert_value_error(
        lambda: validate_selected_recovery_source_record(
            {**jobs_record.to_dict(), "selection_status": "approved_evidence"},
            candidates,
        ),
        "selection_status is unsupported",
    )
    assert_value_error(
        lambda: validate_selected_recovery_source_record(
            {**jobs_record.to_dict(), "selected_source_url": ""},
            candidates,
        ),
        "selected_source_url must be a non-empty string",
    )
    assert_value_error(
        lambda: validate_selected_recovery_source_record(
            {**jobs_record.to_dict(), "approved_evidence": True},
            candidates,
        ),
        "approved_evidence must be false",
    )
    assert_value_error(
        lambda: validate_selected_recovery_source_record(
            {**jobs_record.to_dict(), "public_ready": True},
            candidates,
        ),
        "public_ready must be false",
    )
    assert_value_error(
        lambda: validate_selected_recovery_source_record(
            {
                **jobs_record.to_dict(),
                "selected_recovery_candidate_id": "RECOVERY_UNKNOWN",
            },
            candidates,
        ),
        "selected_recovery_candidate_id is unknown",
    )

    protected_paths = [
        Path("data/evidence_seed/videos.yaml"),
        Path("data/real_evidence_inputs/VID_JOBS_001.template.json"),
        Path("data/real_evidence_inputs/VID_HEALTH_001.template.json"),
    ]
    before = {path: path.read_text(encoding="utf-8") for path in protected_paths}
    for command in (
        "python scripts/inspect_selected_recovery_sources.py",
        "python scripts/inspect_selected_recovery_sources.py --evidence-id VID_JOBS_001",
        "python scripts/inspect_selected_recovery_sources.py --evidence-id VID_HEALTH_001",
    ):
        exit_code, output = run_command(command)
        if exit_code != 0:
            raise AssertionError(
                f"Selected recovery source inspection command failed: {command}"
            )
        if "approved_evidence: 0" not in output:
            raise AssertionError("Selected recovery inspection must not approve evidence")
        if "public_ready: False" not in output:
            raise AssertionError("Selected recovery inspection must not mark readiness")
    after = {path: path.read_text(encoding="utf-8") for path in protected_paths}
    if before != after:
        raise AssertionError(
            "Selected recovery source lane must not modify protected source files"
        )

    with tempfile.TemporaryDirectory() as temp_dir:
        payload = json.loads(
            DEFAULT_SELECTED_RECOVERY_SOURCE_PATH.read_text(encoding="utf-8")
        )
        duplicate = json.loads(json.dumps(payload))
        duplicate["selected_sources"].append(dict(duplicate["selected_sources"][0]))
        duplicate_path = Path(temp_dir) / "selected_recovery_sources.json"
        duplicate_path.write_text(json.dumps(duplicate, indent=2), encoding="utf-8")
        assert_value_error(
            lambda: load_selected_recovery_sources(
                path=duplicate_path,
                candidates=candidates,
            ),
            "Duplicate selected recovery source",
        )

        unknown = json.loads(json.dumps(payload))
        unknown["selected_sources"][0]["selected_recovery_candidate_id"] = (
            "RECOVERY_UNKNOWN"
        )
        unknown_path = Path(temp_dir) / "selected_recovery_sources_unknown.json"
        unknown_path.write_text(json.dumps(unknown, indent=2), encoding="utf-8")
        assert_value_error(
            lambda: load_selected_recovery_sources(
                path=unknown_path,
                candidates=candidates,
            ),
            "selected_recovery_candidate_id is unknown",
        )

    print("✓ Selected Recovery Source v1 lane validation OK")


def validate_source_content_extraction_lane() -> None:
    selected_sources = load_selected_recovery_sources()
    if len(selected_sources) != 2:
        raise AssertionError("Source Content Extraction v1 must load 2 sources")

    no_network_record = extract_content_for_selected_source(
        selected_sources[0],
        no_network=True,
        timeout_seconds=1,
    )
    validate_source_content_extraction_record(no_network_record)
    if no_network_record.extraction_status != "not_checked":
        raise AssertionError("No-network extraction must remain not_checked")
    if no_network_record.content_source_status != "not_checked":
        raise AssertionError("No-network extraction must not check source status")
    if not no_network_record.requires_manual_review:
        raise AssertionError("Extraction records must require manual review")
    if (
        no_network_record.approved_evidence
        or no_network_record.public_ready
        or no_network_record.institutional_ready
        or no_network_record.report_ready
    ):
        raise AssertionError("Extraction records must not mark readiness")
    if "verified_for_approval_review" in no_network_record.to_dict():
        raise AssertionError("Extraction records must not include approval verification")
    if "verified_for_manual_review" in no_network_record.to_dict():
        raise AssertionError("Extraction records must not include manual verification")

    jobs_text, jobs_quote = extract_candidate_text(
        "VID_JOBS_001",
        """
        <html><head><script>var ignored = 'jobs';</script></head>
        <body><p>The manifesto commits to create 500,000 jobs through
        employment programmes.</p></body></html>
        """,
    )
    if "500,000 jobs" not in jobs_text:
        raise AssertionError("Jobs extraction fixture did not find candidate text")
    if "500,000 jobs" not in jobs_quote:
        raise AssertionError("Jobs extraction fixture did not find quote candidate")

    health_text, health_quote = extract_candidate_text(
        "VID_HEALTH_001",
        """
        <html><body><p>Botswana declared a public health emergency after
        clinics reported medicine shortage pressure.</p></body></html>
        """,
    )
    if "public health emergency" not in health_text:
        raise AssertionError("Health extraction fixture did not find candidate text")
    if "public health emergency" not in health_quote:
        raise AssertionError("Health extraction fixture did not find quote candidate")

    assert_value_error(
        lambda: validate_source_content_extraction_record(
            {**no_network_record.to_dict(), "source_url": ""}
        ),
        "source_url must be a non-empty string",
    )
    assert_value_error(
        lambda: validate_source_content_extraction_record(
            {**no_network_record.to_dict(), "approved_evidence": True}
        ),
        "approved_evidence must be false",
    )
    assert_value_error(
        lambda: validate_source_content_extraction_record(
            {**no_network_record.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )
    assert_value_error(
        lambda: validate_source_content_extraction_record(
            {**no_network_record.to_dict(), "verified_for_approval_review": False}
        ),
        "must not contain verification field",
    )

    candidate_record = SourceContentExtractionRecord(
        original_evidence_id="VID_JOBS_001",
        selected_recovery_candidate_id="RECOVERY_VID_JOBS_001_UDC_HOME",
        source_url="https://www.udc.org.bw/",
        extraction_status="extracted_candidate",
        content_source_status="source_reachable",
        extracted_text_candidate="The manifesto references 500,000 jobs.",
        extracted_quote_candidate="The manifesto references 500,000 jobs.",
        extraction_method="fixture",
        requires_manual_review=True,
        approved_evidence=False,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        notes="Fixture candidate only. Human review required.",
    )
    validate_source_content_extraction_record(candidate_record)

    protected_paths = [
        Path("data/evidence_seed/videos.yaml"),
        Path("data/real_evidence_inputs/VID_JOBS_001.template.json"),
        Path("data/real_evidence_inputs/VID_HEALTH_001.template.json"),
    ]
    before = {path: path.read_text(encoding="utf-8") for path in protected_paths}
    summary = extract_selected_source_content(no_network=True, timeout_seconds=1)
    after = {path: path.read_text(encoding="utf-8") for path in protected_paths}
    if before != after:
        raise AssertionError(
            "Source content extraction must not modify protected source files"
        )
    if summary["total_sources"] != 2:
        raise AssertionError("Source content extraction no-network must process 2 records")
    if summary["not_checked_count"] != 2:
        raise AssertionError("Source content extraction no-network must not check URLs")
    if summary["requires_manual_review_count"] != 2:
        raise AssertionError("Source content extraction must require manual review")
    if summary["approved_evidence"] != 0:
        raise AssertionError("Source content extraction must not approve evidence")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Source content extraction must not mark readiness")

    for command in (
        "python scripts/extract_selected_source_content.py --no-network",
        "python scripts/extract_selected_source_content.py --evidence-id VID_JOBS_001 --no-network",
        "python scripts/extract_selected_source_content.py --evidence-id VID_HEALTH_001 --no-network",
    ):
        exit_code, output = run_command(command)
        if exit_code != 0:
            raise AssertionError(
                f"Source content extraction command failed: {command}"
            )
        if "approved_evidence: 0" not in output:
            raise AssertionError("Source extraction CLI must not approve evidence")
        if "public_ready: False" not in output:
            raise AssertionError("Source extraction CLI must not mark readiness")

    print("✓ Source Content Extraction v1 lane validation OK")


def validate_source_content_verification_lane() -> None:
    missing_records = build_missing_artifact_records()
    if len(missing_records) != 2:
        raise AssertionError("Missing extraction artifact fallback must cover 2 sources")
    for record in missing_records:
        validate_source_content_verification_record(record)
        if record.verification_status != "not_checked":
            raise AssertionError("Missing extraction artifact must remain not_checked")
        if record.content_review_status != "insufficient_content":
            raise AssertionError("Missing extraction artifact must be insufficient")
        if (
            record.verified_for_content_review
            or record.approved_evidence
            or record.public_ready
            or record.institutional_ready
            or record.report_ready
        ):
            raise AssertionError("Missing extraction artifact must not mark readiness")

    jobs_extraction = SourceContentExtractionRecord(
        original_evidence_id="VID_JOBS_001",
        selected_recovery_candidate_id="RECOVERY_VID_JOBS_001_UDC_HOME",
        source_url="https://www.udc.org.bw/",
        extraction_status="extracted_candidate",
        content_source_status="source_reachable",
        extracted_text_candidate=(
            "The manifesto commits to create 500,000 jobs through employment "
            "programmes."
        ),
        extracted_quote_candidate="The manifesto commits to create 500,000 jobs.",
        extraction_method="fixture",
        requires_manual_review=True,
        approved_evidence=False,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        notes="Fixture candidate only. Human review required.",
    )
    jobs_verification = verify_extraction_record(jobs_extraction)
    validate_source_content_verification_record(jobs_verification)
    if (
        jobs_verification.verification_status
        != "verified_candidate_for_content_review"
    ):
        raise AssertionError("Jobs candidate must verify for content review")
    if jobs_verification.content_review_status != "keyword_match":
        raise AssertionError("Jobs candidate must have keyword_match review status")
    if not jobs_verification.verified_for_content_review:
        raise AssertionError("Jobs candidate must set content-review candidate flag")
    if (
        jobs_verification.approved_evidence
        or jobs_verification.public_ready
        or jobs_verification.institutional_ready
        or jobs_verification.report_ready
    ):
        raise AssertionError("Content-review candidate must not imply approval")

    health_extraction = SourceContentExtractionRecord(
        original_evidence_id="VID_HEALTH_001",
        selected_recovery_candidate_id="RECOVERY_VID_HEALTH_001_REUTERS",
        source_url=(
            "https://www.reuters.com/business/healthcare-pharmaceuticals/"
            "botswana-declares-public-health-emergency-clinics-run-out-medicine-"
            "2025-08-25/"
        ),
        extraction_status="blocked",
        content_source_status="source_blocked",
        extracted_text_candidate="",
        extracted_quote_candidate="",
        extraction_method="fixture",
        requires_manual_review=True,
        approved_evidence=False,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        notes="Fixture blocked source. Fallback review required.",
    )
    health_verification = verify_extraction_record(health_extraction)
    validate_source_content_verification_record(health_verification)
    if health_verification.verification_status != "blocked_pending_fallback":
        raise AssertionError("Blocked health source must remain pending fallback")
    if health_verification.content_review_status != "requires_fallback_source":
        raise AssertionError("Blocked health source must require fallback")
    if health_verification.verified_for_content_review:
        raise AssertionError("Blocked health source must not verify for review")

    assert_value_error(
        lambda: validate_source_content_verification_record(
            {**jobs_verification.to_dict(), "verified_for_approval_review": True}
        ),
        "must not contain verified_for_approval_review",
    )
    assert_value_error(
        lambda: validate_source_content_verification_record(
            {**jobs_verification.to_dict(), "approved_evidence": True}
        ),
        "approved_evidence must be false",
    )
    assert_value_error(
        lambda: validate_source_content_verification_record(
            {**jobs_verification.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )
    assert_value_error(
        lambda: validate_source_content_verification_record(
            {**jobs_verification.to_dict(), "source_url": ""}
        ),
        "source_url must be a non-empty string",
    )

    protected_paths = [
        Path("data/evidence_seed/videos.yaml"),
        Path("data/real_evidence_inputs/VID_JOBS_001.template.json"),
        Path("data/real_evidence_inputs/VID_HEALTH_001.template.json"),
    ]
    before = {path: path.read_text(encoding="utf-8") for path in protected_paths}

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        missing_summary = verify_source_content(
            extraction_status_path=temp_path / "missing_status.json",
            output_dir=temp_path / "missing_output",
            no_network=True,
        )
        if missing_summary["artifact_found"]:
            raise AssertionError("Missing extraction artifact must report artifact_found=false")
        if missing_summary["not_checked_count"] != 2:
            raise AssertionError("Missing extraction artifact must keep records unchecked")

        status_path = temp_path / "source_content_extraction_status.json"
        status_path.write_text(
            json.dumps(
                {
                    "records": [
                        jobs_extraction.to_dict(),
                        health_extraction.to_dict(),
                    ]
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        summary = verify_source_content(
            extraction_status_path=status_path,
            output_dir=temp_path / "verification_output",
            no_network=True,
        )
        if summary["total_records"] != 2:
            raise AssertionError("Source content verification must process 2 records")
        if summary["verified_candidate_for_content_review_count"] != 1:
            raise AssertionError("Jobs source must be one content-review candidate")
        if summary["blocked_pending_fallback_count"] != 1:
            raise AssertionError("Health source must be one blocked fallback")
        if summary["approved_evidence"] != 0:
            raise AssertionError("Source content verification must not approve evidence")
        if (
            summary["public_ready"]
            or summary["institutional_ready"]
            or summary["report_ready"]
        ):
            raise AssertionError("Source content verification must not mark readiness")

        jobs_summary = verify_source_content(
            evidence_id="VID_JOBS_001",
            extraction_status_path=status_path,
            output_dir=temp_path / "jobs_verification_output",
            no_network=True,
        )
        if jobs_summary["verified_candidate_for_content_review_count"] != 1:
            raise AssertionError("Source verification evidence-id filter failed")

    summary = verify_source_content(no_network=True)
    after = {path: path.read_text(encoding="utf-8") for path in protected_paths}
    if before != after:
        raise AssertionError(
            "Source content verification must not modify protected source files"
        )
    if summary["approved_evidence"] != 0:
        raise AssertionError("Source content verification must not approve evidence")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Source content verification must not mark readiness")

    for command in (
        "python scripts/verify_source_content.py --no-network",
        "python scripts/verify_source_content.py --evidence-id VID_JOBS_001 --no-network",
        "python scripts/verify_source_content.py --evidence-id VID_HEALTH_001 --no-network",
    ):
        exit_code, output = run_command(command)
        if exit_code != 0:
            raise AssertionError(
                f"Source content verification command failed: {command}"
            )
        if "approved_evidence: 0" not in output:
            raise AssertionError("Source verification CLI must not approve evidence")
        if "public_ready: False" not in output:
            raise AssertionError("Source verification CLI must not mark readiness")

    print("✓ Source Content Verification v1 lane validation OK")


def validate_real_evidence_approval_lane() -> None:
    blocked = RealEvidenceApprovalRecord(
        approval_id="APPROVAL_TEST_BLOCKED",
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        source_url="https://www.youtube.com/watch?v=e0MLzB5nGDc",
        transcript_status="real_transcript_approval_required",
        timestamp_status="real_timestamp_approval_required",
        quote_status="real_quote_approval_required",
        context_status="context_approval_required",
        case_relevance_status="case_relevance_review_required",
        reviewer="real-evidence-approval-v1-test",
        approval_status="blocked_manual_review_required",
        approval_candidate=False,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        approval_notes=(
            "Blocked test record. Real transcript, timestamp, quote, context, "
            "case relevance, and reviewer approval are required."
        ),
        blocker_reasons=list(APPROVAL_DRY_RUN_BLOCKER_REASONS),
    )
    validate_real_evidence_approval_record(blocked)

    approved_candidate = RealEvidenceApprovalRecord(
        approval_id="APPROVAL_TEST_CANDIDATE",
        evidence_id="VID_JOBS_001",
        case_id="CASE_002",
        source_url="https://www.youtube.com/watch?v=e0MLzB5nGDc",
        transcript_status="real_transcript_verified",
        timestamp_status="real_timestamp_verified",
        quote_status="real_quote_verified",
        context_status="context_verified",
        case_relevance_status="case_relevance_verified",
        reviewer="real-evidence-approval-v1-test",
        approval_status="approved_evidence_candidate",
        approval_candidate=True,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        approval_notes=(
            "Candidate test record only. Candidate approval remains below "
            "public, institutional, and report readiness."
        ),
        blocker_reasons=[],
    )
    validate_real_evidence_approval_record(approved_candidate)

    for field_name in (
        "approval_id",
        "evidence_id",
        "case_id",
        "source_url",
        "reviewer",
        "approval_status",
        "approval_notes",
    ):
        assert_value_error(
            lambda field_name=field_name: validate_real_evidence_approval_record(
                {**blocked.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {**blocked.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {**blocked.to_dict(), "institutional_ready": True}
        ),
        "institutional_ready must be false",
    )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {**blocked.to_dict(), "report_ready": True}
        ),
        "report_ready must be false",
    )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {
                **approved_candidate.to_dict(),
                "transcript_status": "real_transcript_approval_required",
            }
        ),
        "approval_candidate requires transcript_status=real_transcript_verified",
    )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {
                **approved_candidate.to_dict(),
                "approval_status": "blocked_manual_review_required",
                "blocker_reasons": list(APPROVAL_DRY_RUN_BLOCKER_REASONS),
            }
        ),
        "approval_candidate requires approval_status=approved_evidence_candidate",
    )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {**blocked.to_dict(), "blocker_reasons": []}
        ),
        "blocked statuses require blocker_reasons",
    )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {
                **blocked.to_dict(),
                "approval_status": "rejected_evidence",
                "blocker_reasons": [],
            }
        ),
        "rejected_evidence requires blocker_reasons",
    )

    assert_value_error(
        lambda: validate_real_evidence_approval_record(
            {**approved_candidate.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )
    if (
        approved_candidate.public_ready
        or approved_candidate.institutional_ready
        or approved_candidate.report_ready
    ):
        raise AssertionError(
            "approved_evidence_candidate must not imply release readiness"
        )

    summary = approve_real_evidence_dry_run()
    if summary["total_evidence_items_evaluated"] != 2:
        raise AssertionError("Real Evidence Approval v1 dry-run must produce 2 records")
    if summary["blocked_manual_review_required"] != 2:
        raise AssertionError("Real Evidence Approval v1 dry-run must block records")
    if summary["approved_evidence_candidate"] != 0:
        raise AssertionError("Real Evidence Approval v1 dry-run must not approve")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Real Evidence Approval v1 must not mark readiness")
    for record in summary["records"]:
        if record.blocker_reasons != APPROVAL_DRY_RUN_BLOCKER_REASONS:
            raise AssertionError("Real Evidence Approval v1 blockers must be deterministic")

    policy_summary = check_release_policy_dry_run()
    if policy_summary["manual_override_allowed"] is not True:
        raise AssertionError("Release Policy v1 must still work")

    readiness_summary = check_release_readiness_dry_run()
    if readiness_summary["release_blocked"] is not True:
        raise AssertionError("Release Readiness v1 must remain blocked")

    assert_nonempty_file(str(DEFAULT_REAL_EVIDENCE_APPROVAL_RECORDS_OUTPUT))
    assert_nonempty_file(str(DEFAULT_REAL_EVIDENCE_APPROVAL_SUMMARY_OUTPUT))
    print("✓ Real Evidence Approval v1 lane validation OK")


def validate_final_approved_packet_lane() -> None:
    blocked = FinalApprovedEvidencePacketRecord(
        packet_id="FINAL_APPROVED_EVIDENCE_PACKET_TEST_BLOCKED",
        report_id="DUMA_BOKO_REVIEW_DRAFT_V1_FIXTURE",
        packet_status="blocked_no_approved_evidence",
        approved_evidence_count=0,
        blocked_evidence_count=2,
        approved_evidence_ids=[],
        blocked_evidence_ids=["VID_JOBS_001", "VID_HEALTH_001"],
        source_table_status="blocked_pending_approval",
        transcript_table_status="blocked_pending_approval",
        timestamp_table_status="blocked_pending_approval",
        quote_table_status="blocked_pending_approval",
        manual_review_table_status="available_but_not_release_approval",
        approval_table_status="no_approved_evidence_candidates",
        release_readiness_status="blocked_manual_review_required",
        release_policy_status="active_local_ci_override_policy",
        audit_trail_status="dry_run_audit_available",
        packet_ready=False,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        blocker_reasons=list(PACKET_DRY_RUN_BLOCKER_REASONS),
        packet_notes=(
            "Blocked packet test record. No approved evidence candidates are "
            "available for packet readiness."
        ),
    )
    validate_final_approved_packet_record(blocked)

    approved_candidate = FinalApprovedEvidencePacketRecord(
        packet_id="FINAL_APPROVED_EVIDENCE_PACKET_TEST_CANDIDATE",
        report_id="DUMA_BOKO_REVIEW_DRAFT_V1_FIXTURE",
        packet_status="approved_packet_candidate",
        approved_evidence_count=1,
        blocked_evidence_count=0,
        approved_evidence_ids=["VID_JOBS_001"],
        blocked_evidence_ids=[],
        source_table_status="available_from_approved_candidates",
        transcript_table_status="available_from_approved_candidates",
        timestamp_table_status="available_from_approved_candidates",
        quote_table_status="available_from_approved_candidates",
        manual_review_table_status="available_from_approved_candidates",
        approval_table_status="approved_evidence_candidates_available",
        release_readiness_status="release_candidate_pending_approval",
        release_policy_status="active_local_ci_override_policy",
        audit_trail_status="dry_run_audit_available",
        packet_ready=True,
        public_ready=False,
        institutional_ready=False,
        report_ready=False,
        blocker_reasons=[],
        packet_notes=(
            "Approved packet candidate test record only. Packet candidacy remains "
            "below public, institutional, and report readiness."
        ),
    )
    validate_final_approved_packet_record(approved_candidate)

    for field_name in ("packet_id", "report_id", "packet_status", "packet_notes"):
        assert_value_error(
            lambda field_name=field_name: validate_final_approved_packet_record(
                {**blocked.to_dict(), field_name: ""}
            ),
            f"{field_name} must be a non-empty string",
        )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {**blocked.to_dict(), "packet_ready": True}
        ),
        "packet_ready requires approved_evidence_count > 0",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {**blocked.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {**blocked.to_dict(), "institutional_ready": True}
        ),
        "institutional_ready must be false",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {**blocked.to_dict(), "report_ready": True}
        ),
        "report_ready must be false",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {**blocked.to_dict(), "blocker_reasons": []}
        ),
        "blocked packet statuses require blocker_reasons",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {
                **blocked.to_dict(),
                "packet_status": "blocked_release_not_authorized",
                "blocker_reasons": [],
            }
        ),
        "blocked packet statuses require blocker_reasons",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {
                **approved_candidate.to_dict(),
                "approved_evidence_count": 0,
                "approved_evidence_ids": [],
                "packet_ready": False,
            }
        ),
        "approved_packet_candidate requires approved_evidence_count > 0",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {
                **approved_candidate.to_dict(),
                "blocked_evidence_count": 1,
                "blocked_evidence_ids": ["VID_JOBS_001"],
            }
        ),
        "blocked evidence IDs must not appear in approved_evidence_ids",
    )

    assert_value_error(
        lambda: validate_final_approved_packet_record(
            {**approved_candidate.to_dict(), "public_ready": True}
        ),
        "public_ready must be false",
    )
    if (
        approved_candidate.public_ready
        or approved_candidate.institutional_ready
        or approved_candidate.report_ready
    ):
        raise AssertionError("approved_packet_candidate must not imply readiness")

    summary = generate_final_approved_packet_dry_run()
    if summary["packet_status"] != "blocked_no_approved_evidence":
        raise AssertionError("Final Approved Evidence Packet v1 must block dry-run")
    if summary["approved_evidence_count"] != 0:
        raise AssertionError("Final Approved Evidence Packet v1 must not approve")
    if summary["blocked_evidence_count"] != 2:
        raise AssertionError("Final Approved Evidence Packet v1 must block 2 records")
    if summary["packet_ready"] is not False:
        raise AssertionError("Final Approved Evidence Packet v1 must not be ready")
    if (
        summary["public_ready"]
        or summary["institutional_ready"]
        or summary["report_ready"]
    ):
        raise AssertionError("Final Approved Evidence Packet v1 must not mark readiness")
    record = summary["record"]
    if record.approved_evidence_ids:
        raise AssertionError("Final Approved Evidence Packet v1 must not approve IDs")
    if record.blocked_evidence_ids != ["VID_JOBS_001", "VID_HEALTH_001"]:
        raise AssertionError("Final Approved Evidence Packet v1 blocked IDs changed")
    if record.blocker_reasons != PACKET_DRY_RUN_BLOCKER_REASONS:
        raise AssertionError("Final Approved Evidence Packet v1 blockers changed")

    approval_summary = approve_real_evidence_dry_run()
    if approval_summary["approved_evidence_candidate"] != 0:
        raise AssertionError("Real Evidence Approval v1 must still block fixtures")
    if approval_summary["blocked_manual_review_required"] != 2:
        raise AssertionError("Real Evidence Approval v1 must still block both records")

    policy_summary = check_release_policy_dry_run()
    if policy_summary["manual_override_allowed"] is not True:
        raise AssertionError("Release Policy v1 must still work")

    readiness_summary = check_release_readiness_dry_run()
    if readiness_summary["release_blocked"] is not True:
        raise AssertionError("Release Readiness v1 must remain blocked")

    assert_nonempty_file(str(DEFAULT_FINAL_APPROVED_PACKET_RECORD))
    assert_nonempty_file(str(DEFAULT_FINAL_APPROVED_PACKET_SUMMARY))
    print("✓ Final Approved Evidence Packet v1 lane validation OK")


def validate_seed_evidence_index() -> None:
    seed_evidence = load_seed_evidence()
    if len(seed_evidence) != 2:
        raise AssertionError(f"Expected 2 seed evidence records, got {len(seed_evidence)}")

    index = build_evidence_index()
    if index["metadata"]["total_evidence_records"] != 2:
        raise AssertionError("evidence_index.json must contain 2 evidence records")

    assert_nonempty_file(str(DEFAULT_INDEX_PATH))
    print("✓ seed evidence index OK: 2 records")


def validate_report_source_text() -> None:
    forbidden_text = " ".join(("VALIDATED", "VIDEO", "EVIDENCE", "ATTACHED"))
    for path in (
        Path("evidence/word_exporter.py"),
        Path("evidence/final_report_generator.py"),
        Path("run_final_report.py"),
    ):
        if forbidden_text in path.read_text(encoding="utf-8"):
            raise AssertionError(f"{path} contains overclaimed report text")

    print("✓ report source text avoids overclaimed validation wording")


def main() -> int:
    validate_schema_objects()
    validate_ingestion_lane()
    validate_transcript_acquisition_lane()
    validate_timestamp_verification_lane()
    validate_quote_verification_lane()
    validate_case_evidence_linking_lane()
    validate_report_section_assembly_lane()
    validate_final_report_generation_lane()
    validate_real_evidence_replacement_lane()
    validate_manual_review_lane()
    validate_final_report_hardening_lane()
    validate_release_readiness_lane()
    validate_release_policy_lane()
    validate_real_evidence_inputs_lane()
    validate_real_evidence_auto_collection_helper_lane()
    validate_source_recovery_lane()
    validate_recovery_candidate_verification_lane()
    validate_selected_recovery_source_lane()
    validate_source_content_extraction_lane()
    validate_source_content_verification_lane()
    validate_real_evidence_approval_lane()
    validate_final_approved_packet_lane()
    validate_gate_rejections()
    validate_seed_evidence_index()
    validate_report_source_text()

    run_divergence_engine(
        Path("config/contradiction_targets.yaml"),
        Path("outputs/cases/divergence_cases.json"),
    )

    divergence_json = assert_nonempty_file("outputs/cases/divergence_cases.json")
    validate_divergence_cases(divergence_json)

    exporter = WordExporter(Path("outputs/reports"))
    assert_value_error(
        lambda: exporter.generate_report(
            json.loads(divergence_json.read_text(encoding="utf-8")).get("cases", []),
            "DUMA_BOKO_DIVERGENCE_REPORT_BLOCKED.docx",
        ),
        "EvidenceObject.url must be a non-empty string",
    )

    rc, output = run_command("python run_final_report.py")
    if rc == 0:
        print("run_final_report.py unexpectedly succeeded")
        return 1
    if "EvidenceObject.url must be a non-empty string" not in output:
        print("run_final_report.py failed for an unexpected reason")
        return 1

    print("=" * 70)
    print("ALL v3.0 EVIDENCE GATE TESTS PASSED")
    print("=" * 70)
    return 0


if __name__ == "__main__":
    sys.exit(main())
